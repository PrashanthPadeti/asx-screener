"""
Generate the ASX Screener End-to-End Data Flow Architecture Word document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"C:\Users\Dell\My Claude\Code\ASX Screener\ASX_Screener_Data_Flow_Architecture.docx"

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helper: set paragraph font ────────────────────────────────────────────────
def set_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1, color=(0, 70, 127)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = 'Calibri'
    return p

def add_para(text='', bold=False, size=11, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def add_code_block(lines):
    """Add a monospaced code/diagram block."""
    if isinstance(lines, str):
        lines = lines.split('\n')
    style = doc.styles['Normal']
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name  = 'Courier New'
        run.font.size  = Pt(8.5)
        run.font.color.rgb = RGBColor(30, 30, 30)
        # Light grey background via shading
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F4F4F4')
        rPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    # Header row
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.size  = Pt(9.5)
        run.font.name  = 'Calibri'
        run.font.color.rgb = RGBColor(255, 255, 255)
        # Blue fill for header
        tc   = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '003E6B')
        tcPr.append(shd)
    # Data rows
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            run = cells[i].paragraphs[0].runs[0] if cells[i].paragraphs[0].runs else cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()  # spacing after table

def add_separator():
    p = doc.add_paragraph('─' * 110)
    for run in p.runs:
        run.font.name  = 'Courier New'
        run.font.size  = Pt(7)
        run.font.color.rgb = RGBColor(180, 180, 180)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ASX SCREENER')
set_font(run, 'Calibri', 28, bold=True, color=(0, 62, 107))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Complete End-to-End Data Flow Architecture')
set_font(run2, 'Calibri', 16, bold=False, color=(70, 130, 180))

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Covering all stages from source file downloads → staging →\ntransformation layers → universe tables → final metrics tables')
set_font(run3, 'Calibri', 11, color=(80, 80, 80))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Table of Contents', level=1)
toc_entries = [
    ('PART 1',  'High-Level Architecture Diagram'),
    ('PART 2',  'Detailed Data Flow Document'),
    ('Stage 0', 'External Data Sources'),
    ('Stage 1', 'Ingestion & Staging (staging_au.*)'),
    ('Stage 2', 'Canonical Market Tables (market.*)'),
    ('Stage 3A','Daily Metrics Compute (market.daily_metrics)'),
    ('Stage 3B','Periodic Metrics Compute (Weekly / Monthly / Quarterly / Half-Yearly / Yearly)'),
    ('Stage 3C','Fundamental Metrics Compute (market.fundamental_metrics)'),
    ('Stage 3D','Anomaly Detection (market.anomaly_flags)'),
    ('Stage 4', 'Golden Record — Universe Table (screener.universe)'),
    ('Stage 5', 'API & Application Layer'),
    ('',        'Complete Data Lineage Map'),
    ('',        'Scheduling Summary'),
    ('',        'Key Design Principles'),
]
for num, title in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2 if num else 0.5)
    run = p.add_run(f'{num}  {title}' if num else f'      {title}')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = bool(num and 'PART' in num)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 1 — HIGH-LEVEL ARCHITECTURE DIAGRAM
# ══════════════════════════════════════════════════════════════════════════════
add_heading('PART 1 — High-Level Architecture Diagram', level=1)
add_separator()

diagram = r"""
╔══════════════════════════════════════════════════════════════════════════════════════╗
║              ASX SCREENER — DATA FLOW ARCHITECTURE                                   ║
╚══════════════════════════════════════════════════════════════════════════════════════╝

 ┌────────────────────────────────────────────────────────────────────────────────────┐
 │  LAYER 0 — EXTERNAL DATA SOURCES                                                   │
 │                                                                                    │
 │   ┌─────────────┐  ┌─────────────┐  ┌──────────────┐  ┌──────────────┐            │
 │   │  EODHD API  │  │ ASIC Short  │  │ ASX Announce │  │  RBA / ABS   │            │
 │   │  (JSON/CSV) │  │  Positions  │  │  API / Feed  │  │  Macro Data  │            │
 │   └──────┬──────┘  └──────┬──────┘  └──────┬───────┘  └──────┬───────┘            │
 └──────────┼────────────────┼────────────────┼─────────────────┼──────────────────  ┘
            │                │                │                 │
            ▼                ▼                ▼                 ▼
 ┌────────────────────────────────────────────────────────────────────────────────────┐
 │  LAYER 1 — INGESTION / STAGING  (schema: staging_au.*)                                │
 │  TRUNCATE → RELOAD on every run                                                    │
 │                                                                                    │
 │  staging_au.eod_prices   staging_au.fundamentals   staging_au.dividends                    │
 │  staging_au.splits       staging_au.exchange_syms  staging_au.bulk_eod                     │
 │  staging_au.earnings     staging_au.macro_indics   staging_au.short_positions              │
 │  staging_au.announcements                                                             │
 │                                                                                    │
 │  Scripts: daily_pipeline.py / weekly_pipeline.py / monthly_pipeline.py            │
 └────────────────────────────────────────────────────────────────────────────────────┘
            │
            ▼
 ┌────────────────────────────────────────────────────────────────────────────────────┐
 │  LAYER 2 — CANONICAL PRICE & REFERENCE TABLES  (schema: market.*)                 │
 │                                                                                    │
 │  market.securities     market.daily_prices (TimescaleDB)  market.dividends         │
 │  market.fundamentals   market.announcements               market.short_positions   │
 │  market.splits         market.earnings                    market.macro_indicators  │
 └────────────────────────────────────────────────────────────────────────────────────┘
            │
            ▼
 ┌────────────────────────────────────────────────────────────────────────────────────┐
 │  LAYER 3 — COMPUTE ENGINE  (6-tier metrics hierarchy)                              │
 │                                                                                    │
 │  market.daily_prices                                                               │
 │    ├──► market.daily_metrics      (returns, SMA, EMA, RSI, MACD, BB, ATR, vol)    │
 │    ├──► market.weekly_metrics     (OHLCV aggregates, weekly return)                │
 │    ├──► market.monthly_metrics    (monthly OHLCV, monthly return)                  │
 │    ├──► market.quarterly_metrics  (quarterly OHLCV, quarterly return)              │
 │    ├──► market.halfyearly_metrics (6-month OHLCV, 6m return)                      │
 │    └──► market.yearly_metrics     (annual OHLCV, 52w high/low)                    │
 │                                                                                    │
 │  market.daily_prices + fundamentals + dividends                                   │
 │    └──► market.fundamental_metrics (PE, PB, EV/EBITDA, FCF yield, DY, franking)  │
 │                                                                                    │
 │  market.daily_metrics + fundamental_metrics                                        │
 │    └──► market.anomaly_flags  (7 flags: vol spike, gap, reversal, breakout,       │
 │                                 RSI extreme, trend break, new high)               │
 └────────────────────────────────────────────────────────────────────────────────────┘
            │
            ▼
 ┌────────────────────────────────────────────────────────────────────────────────────┐
 │  LAYER 4 — GOLDEN RECORD / UNIVERSE TABLE                                          │
 │                                                                                    │
 │  screener.universe — 458+ column denormalized row per ASX stock                   │
 │    Sources joined: securities + daily_prices + daily_metrics +                    │
 │    weekly/monthly/quarterly/halfyearly/yearly_metrics +                           │
 │    fundamental_metrics + dividends + short_positions +                            │
 │    announcements + anomaly_flags                                                   │
 │                                                                                    │
 │  Rebuilt NIGHTLY via build_screener_universe.py (TRUNCATE + full INSERT)          │
 └────────────────────────────────────────────────────────────────────────────────────┘
            │
            ▼
 ┌────────────────────────────────────────────────────────────────────────────────────┐
 │  LAYER 5 — API & APPLICATION                                                       │
 │                                                                                    │
 │  FastAPI Backend                   Next.js Frontend                                │
 │  /screener  → universe             Stock Screener (filter/sort/paginate)           │
 │  /stocks/{t}→ all layers           Stock Detail (price, technicals, divs)         │
 │  /alerts    → triggers             Price Alerts & Portfolio Tracker                │
 │  /admin     → stats                Admin Console                                   │
 └────────────────────────────────────────────────────────────────────────────────────┘

 ┌────────────────────────────────────────────────────────────────────────────────────┐
 │  ORCHESTRATION                                                                     │
 │                                                                                    │
 │  PM2 / System Cron                    APScheduler (in-process, FastAPI)            │
 │  daily_pipeline.py  (after ASX close) alert_checker          (every 1 min)        │
 │  weekly_pipeline.py (Monday morning)  portfolio_updater       (every 5 min)       │
 │  monthly_pipeline.py (1st of month)   anomaly_detector        (every 1 hour)      │
 │                                       announcement_fetcher    (every 2 hours)     │
 │                                       short_position_updater  (daily)             │
 │                                       + 8 more scheduled tasks                    │
 └────────────────────────────────────────────────────────────────────────────────────┘
"""
add_code_block(diagram)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 2 — DETAILED DATA FLOW DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading('PART 2 — Detailed Data Flow Document', level=1)
add_separator()

# ── STAGE 0 ──────────────────────────────────────────────────────────────────
add_heading('Stage 0 — External Data Sources', level=2)
add_table(
    headers=['Source', 'Data Type', 'Format', 'Frequency', 'Key Fields'],
    rows=[
        ['EODHD API',             'EOD prices, bulk prices, fundamentals, dividends, splits, earnings', 'JSON / CSV', 'Daily (prices)\nWeekly (fundamentals)', 'ticker, date, open, high, low, close, volume, adjusted_close'],
        ['ASIC',                  'Short-selling position reports',                                     'CSV',        'Daily',                                   'ticker, short_position_pct, short_position_count'],
        ['ASX Announcements API', 'Corporate announcements, price-sensitive releases',                  'JSON',       'Every 2 hours',                           'ticker, released_at, headline, category, price_sensitive'],
        ['RBA / ABS',             'Macro indicators (cash rate, CPI, GDP)',                             'JSON',       'Monthly',                                 'indicator_code, period, value'],
    ],
    col_widths=[1.3, 2.2, 0.8, 1.2, 2.5],
)

# ── STAGE 1 ──────────────────────────────────────────────────────────────────
add_heading('Stage 1 — Ingestion & Staging (staging_au.*)', level=2)
add_para('Pattern: Each pipeline run TRUNCATES all staging tables first, then RELOADS from source. No incremental merging at this layer — staging is always a clean snapshot of the latest API data.', bold=False)
doc.add_paragraph()

add_heading('Staging Tables', level=3)
add_table(
    headers=['Table', 'Source', 'Populated By', 'Content'],
    rows=[
        ['staging_au.eod_prices',       'EODHD /eod/{ticker}',            'daily_pipeline.py',   'Per-ticker EOD OHLCV + adjusted close'],
        ['staging_au.bulk_eod',         'EODHD /eod/bulk',                'daily_pipeline.py',   'All-market EOD in a single bulk API call'],
        ['staging_au.fundamentals',     'EODHD /fundamentals/{ticker}',   'weekly_pipeline.py',  'Full JSON blob: financials, balance sheet, cashflow, shares outstanding'],
        ['staging_au.dividends',        'EODHD /dividends/{ticker}',      'weekly + daily',      'Dividend date, amount, franking %, type'],
        ['staging_au.splits',           'EODHD /splits/{ticker}',         'weekly_pipeline.py',  'Split ratio, ex-date'],
        ['staging_au.exchange_symbols', 'EODHD /exchange-symbol-list/AU', 'weekly_pipeline.py',  'All ASX tickers with name, type, ISIN'],
        ['staging_au.earnings',         'EODHD /earnings/{ticker}',       'weekly_pipeline.py',  'EPS actual vs estimate, surprise %, date'],
        ['staging_au.macro_indicators', 'EODHD /macro-indicator/AU',      'monthly_pipeline.py', 'GDP, CPI, unemployment, RBA rate'],
        ['staging_au.short_positions',  'ASIC daily CSV',                 'daily_pipeline.py',   'Ticker, short position %, reporting date'],
        ['staging_au.announcements',    'ASX API',                        'APScheduler (2h)',    'Announcement headline, category, price_sensitive flag'],
    ],
    col_widths=[1.7, 1.8, 1.5, 3.0],
)

add_heading('Pipeline Script: daily_pipeline.py', level=3)
add_code_block("""STEP 1: Authenticate EODHD API
STEP 2: Fetch exchange symbols  →  staging_au.exchange_symbols  (TRUNCATE first)
STEP 3: Fetch bulk EOD prices   →  staging_au.bulk_eod          (TRUNCATE first)
STEP 4: For each active ticker:
           Fetch EOD prices     →  staging_au.eod_prices
           Fetch dividends      →  staging_au.dividends
STEP 5: Download ASIC short position CSV  →  staging_au.short_positions
STEP 6: Promote staging  →  market.* canonical tables  (UPSERT)
STEP 7: Run compute engine (all 6 metric tiers)
STEP 8: Run build_screener_universe.py  (rebuild Golden Record)""")

add_heading('Pipeline Script: weekly_pipeline.py', level=3)
add_code_block("""STEP 1: Fetch exchange symbols (full refresh)
STEP 2: For each ticker:
           Fetch fundamentals  →  staging_au.fundamentals
           Fetch splits        →  staging_au.splits
           Fetch earnings      →  staging_au.earnings
STEP 3: Promote to market.fundamentals, market.splits, market.earnings
STEP 4: Run fundamental_metrics compute engine
STEP 5: Rebuild screener.universe""")

add_heading('Pipeline Script: monthly_pipeline.py', level=3)
add_code_block("""STEP 1: Fetch macro indicators  →  staging_au.macro_indicators
STEP 2: Backfill any missing historical price bars
STEP 3: Recompute full historical metrics for accuracy
STEP 4: Rebuild screener.universe""")

doc.add_page_break()

# ── STAGE 2 ──────────────────────────────────────────────────────────────────
add_heading('Stage 2 — Canonical Market Tables (market.*)', level=2)
add_para('Data is promoted from staging to canonical tables via UPSERT (ON CONFLICT DO UPDATE). These are the source-of-truth normalized tables. The pipeline is idempotent — re-running never creates duplicates.')
doc.add_paragraph()

def add_table_spec(title, rows):
    add_heading(title, level=3)
    add_code_block('\n'.join(rows))

add_table_spec('market.securities — Instrument Master', [
    'Source:  staging_au.exchange_symbols (weekly)',
    'Key:     ticker (ASX code)',
    'Fields:  ticker, isin, name, type (equity/etf/warr), sector, industry,',
    '         gics_sector, gics_group, gics_industry, gics_sub_industry,',
    '         market_cap_band, is_active, listed_date, delisted_date',
    'Flow:    staging_au.exchange_symbols  →  UPSERT  →  market.securities',
])

add_table_spec('market.daily_prices — TimescaleDB Hypertable', [
    'Source:  staging_au.bulk_eod / staging_au.eod_prices (daily)',
    'Key:     (ticker, date)',
    'Fields:  ticker, date, open, high, low, close, adjusted_close, volume,',
    '         adj_factor (split/div adjustment multiplier)',
    'Flow:    staging_au.bulk_eod  →  UPSERT  →  market.daily_prices',
    '         (TimescaleDB chunks by month for query performance)',
    'Note:    adjusted_close = close × adj_factor, accounts for splits and dividends',
])

add_table_spec('market.fundamentals', [
    'Source:  staging_au.fundamentals (weekly)',
    'Key:     (ticker, report_date, period_type)',
    'Fields:  revenue, net_income, ebitda, ebit, gross_profit, operating_income,',
    '         total_assets, total_equity, total_debt, cash_and_equivalents,',
    '         operating_cashflow, free_cashflow, capex,',
    '         shares_outstanding, shares_float, eps_diluted, eps_basic,',
    '         book_value_per_share, report_date, period_type (annual/quarterly)',
    'Flow:    staging_au.fundamentals  →  JSON parse  →  UPSERT  →  market.fundamentals',
])

add_table_spec('market.dividends', [
    'Source:  staging_au.dividends (daily + weekly)',
    'Key:     (ticker, ex_date, payment_type)',
    'Fields:  ticker, ex_date, payment_date, declaration_date,',
    '         amount (AUD), currency, franking_pct (0–100),',
    '         payment_type (cash/stock/special), period (interim/final/special)',
    'Flow:    staging_au.dividends  →  UPSERT  →  market.dividends',
])

add_table_spec('market.short_positions', [
    'Source:  staging_au.short_positions (daily ASIC CSV)',
    'Key:     (ticker, report_date)',
    'Fields:  ticker, report_date, short_position_shares, total_shares,',
    '         short_position_pct = (short_position_shares / total_shares) × 100',
    'Flow:    staging_au.short_positions  →  UPSERT  →  market.short_positions',
])

add_table_spec('market.announcements', [
    'Source:  ASX API (APScheduler, every 2 hours)',
    'Key:     (ticker, announcement_id)',
    'Fields:  ticker, announcement_id, released_at, headline, category,',
    '         is_price_sensitive, document_url, document_size',
    'Flow:    ASX API  →  INSERT IGNORE (append-only, no truncate)  →  market.announcements',
])

doc.add_page_break()

# ── STAGE 3A ─────────────────────────────────────────────────────────────────
add_heading('Stage 3A — Daily Metrics Compute (market.daily_metrics)', level=2)

meta = [
    'Engine:   compute/engine/daily_metrics.py',
    'Source:   market.daily_prices',
    'Target:   market.daily_metrics  (TimescaleDB hypertable, partitioned by month)',
    'Trigger:  After daily_prices UPSERT in daily_pipeline.py',
]
add_code_block(meta)
doc.add_paragraph()

add_heading('Price Return Calculations', level=3)
add_para('All returns use the universal formula:')
add_code_block('  return_Xd = (close[today] - close[today - N trading days]) / close[today - N trading days]')
add_table(
    headers=['Column', 'N (trading days)', 'Business Logic'],
    rows=[
        ['return_1d', '1',   'Yesterday\'s close → today\'s close'],
        ['return_1w', '5',   '~1 calendar week (Mon–Fri)'],
        ['return_1m', '21',  '~1 calendar month'],
        ['return_3m', '63',  '~1 calendar quarter'],
        ['return_6m', '126', '~6 calendar months'],
        ['return_1y', '252', '~1 calendar year (full trading year)'],
    ],
    col_widths=[1.2, 1.5, 5.3],
)

add_heading('Moving Average Calculations', level=3)
add_table(
    headers=['Column', 'Formula', 'Window'],
    rows=[
        ['sma_20',  'AVG(close) OVER (ROWS 19 PRECEDING)',                               '20 trading days'],
        ['sma_50',  'AVG(close) OVER (ROWS 49 PRECEDING)',                               '50 trading days'],
        ['sma_200', 'AVG(close) OVER (ROWS 199 PRECEDING)',                              '200 trading days'],
        ['ema_12',  'close × (2/13) + ema_12_prev × (11/13)',                            'Smoothing = 2/(12+1)'],
        ['ema_26',  'close × (2/27) + ema_26_prev × (25/27)',                            'Smoothing = 2/(26+1)'],
    ],
    col_widths=[1.0, 3.5, 2.0],
)

add_heading('Momentum Indicators', level=3)
add_para('RSI (14-day):', bold=True)
add_code_block("""  gain_avg = AVG of gains over last 14 days (where close > prev_close)
  loss_avg = AVG of losses over last 14 days (where close < prev_close)
  RS       = gain_avg / loss_avg
  RSI      = 100 - (100 / (1 + RS))

  Stored as rsi_14 (range 0–100).  RSI > 70 = overbought;  RSI < 30 = oversold.""")

add_para('MACD:', bold=True)
add_code_block("""  macd_line    = ema_12 - ema_26
  signal_line  = EMA(9) of macd_line
  macd_hist    = macd_line - signal_line""")

add_para('Bollinger Bands (20-day, 2σ):', bold=True)
add_code_block("""  bb_middle = sma_20
  bb_std    = STDDEV(close) OVER (20 days)
  bb_upper  = sma_20 + (2 × bb_std)
  bb_lower  = sma_20 - (2 × bb_std)
  bb_width  = (bb_upper - bb_lower) / bb_middle        -- normalised band width
  bb_pct_b  = (close - bb_lower) / (bb_upper - bb_lower)   -- position within bands (0–1)""")

add_para('ATR (Average True Range, 14-day):', bold=True)
add_code_block("""  true_range = MAX( high - low,
                   ABS(high - prev_close),
                   ABS(low  - prev_close) )
  atr_14     = EMA(14) of true_range""")

add_para('Volume Indicators:', bold=True)
add_code_block("""  volume_sma_20    = AVG(volume) OVER 20 days
  volume_ratio     = volume / volume_sma_20          -- 1.0 = average;  >2.0 = elevated
  volume_change_1d = (volume - prev_volume) / prev_volume""")

add_para('Cross & Position Fields:', bold=True)
add_code_block("""  price_vs_sma_20  = close / sma_20 - 1        -- % above/below 20d SMA
  price_vs_sma_50  = close / sma_50 - 1
  price_vs_sma_200 = close / sma_200 - 1
  above_sma_20     = 1 if close > sma_20  else 0
  above_sma_50     = 1 if close > sma_50  else 0
  above_sma_200    = 1 if close > sma_200 else 0
  golden_cross     = 1 if sma_50 > sma_200 (bullish)
  death_cross      = 1 if sma_50 < sma_200 (bearish)""")

doc.add_page_break()

# ── STAGE 3B ─────────────────────────────────────────────────────────────────
add_heading('Stage 3B — Periodic Metrics Compute (Weekly / Monthly / Quarterly / Half-Yearly / Yearly)', level=2)
add_para('All period metrics follow the same structural pattern: aggregate market.daily_prices into canonical period bars, then compute period-specific statistics.')
doc.add_paragraph()

periods = [
    ('market.weekly_metrics', 'compute/engine/weekly_metrics.py', 'ISO week (Mon–Fri)', '(ticker, week_start_date)', [
        '  open          = FIRST(open) for the week  (Monday open)',
        '  high          = MAX(high) for the week',
        '  low           = MIN(low) for the week',
        '  close         = LAST(close) for the week  (Friday close)',
        '  volume        = SUM(volume) for the week',
        '  weekly_return = (week_close - prev_week_close) / prev_week_close',
        '  week_range_pct= (high - low) / low',
    ]),
    ('market.monthly_metrics', 'compute/engine/monthly_metrics.py', 'Calendar month', '(ticker, year, month)', [
        '  open             = FIRST(open) of month',
        '  high             = MAX(high) of month',
        '  low              = MIN(low) of month',
        '  close            = LAST(close) of month',
        '  volume           = SUM(volume) of month',
        '  avg_daily_vol    = AVG(volume per day) of month',
        '  trading_days     = COUNT(distinct dates)',
        '  monthly_return   = (month_close - prev_month_close) / prev_month_close',
    ]),
    ('market.quarterly_metrics', 'compute/engine/quarterly_metrics.py', 'Calendar quarter (Q1–Q4)', '(ticker, year, quarter)', [
        '  open              = FIRST(open) of quarter',
        '  high/low/close    = MAX/MIN/LAST as above',
        '  quarterly_return  = (quarter_close - prev_quarter_close) / prev_quarter_close',
        '  avg_daily_volume  = SUM(monthly_volume) / total_trading_days',
    ]),
    ('market.halfyearly_metrics', 'compute/engine/halfyearly_metrics.py', 'H1 (Jan–Jun), H2 (Jul–Dec)', '(ticker, year, half)', [
        '  halfyearly_return = (half_close - prev_half_close) / prev_half_close',
        '  6m_range_pct      = (high - low) / low',
    ]),
    ('market.yearly_metrics', 'compute/engine/yearly_metrics.py', 'Calendar year', '(ticker, year)', [
        '  open            = FIRST(open) of year  (first trading day)',
        '  high            = MAX(high) of year',
        '  low             = MIN(low) of year',
        '  close           = LAST(close) of year',
        '  52w_high        = high  (same as yearly high)',
        '  52w_low         = low   (same as yearly low)',
        '  52w_high_pct    = (close - 52w_low) / (52w_high - 52w_low)',
        '  ytd_return      = (close - year_open) / year_open',
        '  yearly_return   = (year_close - prev_year_close) / prev_year_close',
    ]),
]

for table_name, engine, period, key, calcs in periods:
    add_heading(table_name, level=3)
    add_code_block([
        f'Engine:  {engine}',
        f'Source:  market.daily_prices',
        f'Period:  {period}',
        f'Key:     {key}',
        '',
        'Calculations:',
    ] + [f'  {c}' for c in calcs])

doc.add_page_break()

# ── STAGE 3C ─────────────────────────────────────────────────────────────────
add_heading('Stage 3C — Fundamental Metrics Compute (market.fundamental_metrics)', level=2)
add_code_block([
    'Engine:   compute/engine/fundamental_metrics.py',
    'Sources:  market.fundamentals + market.daily_prices + market.dividends + market.securities',
    'Target:   market.fundamental_metrics',
    'Trigger:  weekly_pipeline.py (full recompute) + daily_pipeline.py (price-sensitive fields)',
])
doc.add_paragraph()

add_heading('Valuation Ratios', level=3)
add_code_block("""  Market Cap        = close × shares_outstanding
  Enterprise Value  = market_cap + total_debt - cash_and_equivalents

  PE Ratio          = close / eps_diluted
                      (NULL if eps_diluted ≤ 0)

  Price/Book        = close / book_value_per_share
                      book_value_per_share = total_equity / shares_outstanding

  Price/Sales       = market_cap / revenue
                      (trailing 12-month revenue)

  EV/EBITDA         = enterprise_value / ebitda
                      (NULL if ebitda ≤ 0)

  EV/Revenue        = enterprise_value / revenue
  EV/EBIT           = enterprise_value / ebit""")

add_heading('Profitability & Margin Ratios', level=3)
add_code_block("""  Gross Margin      = gross_profit / revenue
  Operating Margin  = operating_income / revenue
  Net Margin        = net_income / revenue
  EBITDA Margin     = ebitda / revenue

  ROE               = net_income / total_equity            (trailing 12 months)
  ROA               = net_income / total_assets
  ROCE              = ebit / (total_assets - current_liabilities)

  Revenue Growth    = (revenue_current - revenue_prior) / ABS(revenue_prior)
                      (periods: QoQ, YoY — using same period prior year)
  Earnings Growth   = (eps_current - eps_prior) / ABS(eps_prior)""")

add_heading('Cashflow & Balance Sheet Ratios', level=3)
add_code_block("""  FCF               = operating_cashflow - capex
  FCF Yield         = FCF / market_cap
  FCF per Share     = FCF / shares_outstanding

  Debt/Equity       = total_debt / total_equity
  Debt/EBITDA       = total_debt / ebitda
  Net Debt          = total_debt - cash_and_equivalents
  Net Debt/EBITDA   = net_debt / ebitda
  Interest Coverage = ebit / interest_expense
  Current Ratio     = current_assets / current_liabilities
  Quick Ratio       = (current_assets - inventory) / current_liabilities""")

add_heading('Dividend & Franking Credit Calculations', level=3)
add_code_block("""  Dividend Yield     = annual_dividend_amount / close × 100

  Annual Dividend    = SUM of all dividends in trailing 12 months
                       (interim + final + special)

  Franking Credit    = dividend_amount × (franking_pct / 100) × (0.30 / 0.70)
    ↑ Tax credit per share for an Australian taxpayer.
      Formula: amount × franking_ratio × tax_rate / (1 - tax_rate)
      0.30 = Australian corporate tax rate
      0.70 = after-tax retained portion  (1 - 0.30)

  Grossed-Up Yield   = (dividend_amount + franking_credit) / close × 100
                       Effective yield including the franking credit (full-rate taxpayer)

  Payout Ratio       = annual_dividend_amount / eps_diluted × 100
                       >100% = paying out more than earned (potentially unsustainable)""")

add_heading('Franking Credit Example', level=3)
add_code_block("""  Stock pays $0.50 dividend, 70% franked:

    franking_credit      = $0.50 × (70/100) × (0.30/0.70)
                         = $0.50 × 0.70 × 0.4286
                         = $0.15

    Grossed-up dividend  = $0.50 + $0.15  =  $0.65
    Grossed-up yield (price = $10.00)     =  6.5%""")

doc.add_page_break()

# ── STAGE 3D ─────────────────────────────────────────────────────────────────
add_heading('Stage 3D — Anomaly Detection (market.anomaly_flags)', level=2)
add_code_block([
    'Engine:   compute/engine/anomaly_detection.py',
    'Sources:  market.daily_metrics + market.daily_prices + market.fundamental_metrics',
    'Target:   market.anomaly_flags',
    'Trigger:  APScheduler (hourly during market hours) + daily_pipeline.py (post-close)',
])
doc.add_paragraph()
add_para('Seven anomaly flag types are computed:', bold=True)
doc.add_paragraph()

flags = [
    ('Flag 1: Volume Spike', [
        'Condition:  volume > (volume_sma_20 × volume_spike_threshold)',
        '            Default threshold = 3.0  (300% of average)',
        'Severity:   HIGH if volume > 5× average;  MEDIUM if 3–5×',
        'Expires:    After 1 day',
        'Business:   Unusual trading activity — often precedes announcements or price moves',
    ]),
    ('Flag 2: Price Gap', [
        'Condition:  ABS(open - prev_close) / prev_close > gap_threshold',
        '            Default threshold = 0.03  (3% gap at open)',
        'Direction:  gap_up   if open > prev_close × 1.03',
        '            gap_down if open < prev_close × 0.97',
        'Expires:    After 3 days',
        'Business:   Overnight news, earnings, or corporate actions causing gap at open',
    ]),
    ('Flag 3: Reversal Signal', [
        'Condition:  Bullish: return_1d > +5%  AND  prev_return_1d < -3%',
        '            Bearish: return_1d < -5%  AND  prev_return_1d > +3%',
        'Expires:    After 5 days',
        'Business:   Potential trend exhaustion or capitulation events',
    ]),
    ('Flag 4: Breakout', [
        'Condition:  close > MAX(high over last 52 weeks)  →  52-week high breakout',
        '            close < MIN(low  over last 52 weeks)  →  52-week low breakdown',
        'Expires:    After 10 days',
        'Business:   Multi-year technical breakouts — momentum signal for trend traders',
    ]),
    ('Flag 5: RSI Extreme', [
        'Condition:  rsi_14 > 80  →  overbought extreme',
        '            rsi_14 < 20  →  oversold extreme',
        'Expires:    When RSI re-enters 30–70 range',
        'Business:   Tighter than standard 70/30 — only extreme momentum exhaustion readings',
    ]),
    ('Flag 6: Trend Break', [
        'Condition:  Golden Cross: sma_50 crosses ABOVE sma_200',
        '            Death Cross:  sma_50 crosses BELOW sma_200',
        '            (detected by comparing today\'s vs yesterday\'s cross flag)',
        'Expires:    After 20 days',
        'Business:   Long-term trend change signal',
    ]),
    ('Flag 7: New High / Low', [
        'Condition:  close = MAX(close over last N days)  →  new N-day high',
        '            close = MIN(close over last N days)  →  new N-day low',
        'Periods:    20-day, 52-week (252 trading days), All-time',
        'Expires:    After 1 day',
        'Business:   Price discovery — stocks hitting multi-period extremes',
    ]),
]

for flag_title, flag_lines in flags:
    add_heading(flag_title, level=3)
    add_code_block(flag_lines)

doc.add_page_break()

# ── STAGE 4 ──────────────────────────────────────────────────────────────────
add_heading('Stage 4 — Golden Record (screener.universe)', level=2)
add_code_block([
    'Script:   scripts/eodhd/v2/build_screener_universe.py',
    'Target:   screener.universe',
    'Trigger:  Final step of daily_pipeline.py (after all compute layers complete)',
    'Method:   Complete TRUNCATE + full rebuild from all upstream tables (not incremental)',
    'Size:     458+ columns, one row per active ASX stock',
    '',
    'Purpose:  Pre-joined denormalized table used as the primary query target for all',
    '          screener API endpoints — eliminates expensive JOIN chains at query time.',
])
doc.add_paragraph()

add_heading('Column Groups in screener.universe', level=3)
col_groups = [
    ('GROUP 1 — Identity', 'market.securities', 'ticker, isin, name, short_name, listing_date, delisted_date, is_active, stock_type, sector, industry, gics_sector, gics_industry_group, gics_industry, gics_sub_industry'),
    ('GROUP 2 — Latest Price', 'market.daily_prices (latest date)', 'last_close, last_open, last_high, last_low, last_volume, last_adjusted_close, last_trade_date, adj_factor'),
    ('GROUP 3 — Market Cap & Size', 'Computed at build time', 'market_cap, market_cap_usd, market_cap_band (nano/micro/small/mid/large), shares_outstanding, shares_float'),
    ('GROUP 4 — Price Returns', 'market.daily_metrics (latest)', 'return_1d, return_1w, return_1m, return_3m, return_6m, return_1y'),
    ('GROUP 5 — Technical Indicators', 'market.daily_metrics (latest)', 'sma_20/50/200, ema_12/26, macd_line, signal_line, macd_hist, rsi_14, bb_upper/middle/lower/width/pct_b, atr_14, above_sma_20/50/200, price_vs_sma_20/50/200, golden_cross, death_cross, volume_sma_20, volume_ratio'),
    ('GROUP 6 — Weekly/Monthly Aggregates', 'weekly/monthly_metrics (latest)', 'weekly_return, weekly_volume, weekly_high, weekly_low, monthly_return, monthly_volume, monthly_avg_daily_vol'),
    ('GROUP 7 — Quarterly/Yearly', 'quarterly/yearly_metrics (latest)', 'quarterly_return, ytd_return, yearly_return, 52w_high, 52w_low, 52w_high_pct'),
    ('GROUP 8 — Fundamental Metrics', 'market.fundamental_metrics (latest)', 'pe_ratio, pb_ratio, ps_ratio, ev_ebitda, ev_revenue, market_cap, enterprise_value, net_debt, gross_margin, operating_margin, net_margin, ebitda_margin, roe, roa, roce, revenue, ebitda, net_income, free_cashflow, revenue_growth_yoy, earnings_growth_yoy, debt_equity, debt_ebitda, interest_coverage, fcf_yield, fcf_per_share'),
    ('GROUP 9 — Dividends & Franking', 'market.dividends (trailing 12m)', 'dividend_yield, annual_dividend, franking_pct, grossed_up_yield, franking_credit_per_share, last_ex_date, last_payment_date, last_dividend_amount, is_dividend_paying, dividend_frequency'),
    ('GROUP 10 — Short Positions', 'market.short_positions (latest)', 'short_position_pct, short_position_shares, short_report_date'),
    ('GROUP 11 — Anomaly Flags', 'market.anomaly_flags (is_active=TRUE)', 'has_volume_spike, has_price_gap, has_reversal, has_breakout, has_rsi_extreme, has_trend_break, has_new_high, anomaly_count, anomaly_flag_bitmask'),
    ('GROUP 12 — Announcement Activity', 'market.announcements (trailing 30d)', 'announcement_count_7d, announcement_count_30d, price_sensitive_count_30d, last_announcement_date'),
    ('GROUP 13 — User Engagement', 'user tables', 'watchlist_count, alert_count'),
]
add_table(
    headers=['Group', 'Source', 'Key Columns'],
    rows=col_groups,
    col_widths=[1.8, 2.0, 4.2],
)

add_heading('Universe Rebuild SQL Pattern', level=3)
add_code_block("""  TRUNCATE screener.universe;

  INSERT INTO screener.universe
  SELECT
    s.ticker, s.name, s.sector,
    -- ... all 458 columns ...
    p.close            AS last_close,
    dm.return_1d,
    dm.rsi_14,
    fm.pe_ratio,
    dv.dividend_yield,
    dv.franking_pct,
    sp.short_position_pct,
    COUNT(af.id) FILTER (WHERE af.is_active) AS anomaly_count

  FROM market.securities s

  LEFT JOIN LATERAL (
    SELECT * FROM market.daily_prices
    WHERE ticker = s.ticker
    ORDER BY date DESC LIMIT 1
  ) p ON TRUE

  LEFT JOIN LATERAL (
    SELECT * FROM market.daily_metrics
    WHERE ticker = s.ticker
    ORDER BY date DESC LIMIT 1
  ) dm ON TRUE

  LEFT JOIN LATERAL (
    SELECT * FROM market.fundamental_metrics
    WHERE ticker = s.ticker
    ORDER BY computed_at DESC LIMIT 1
  ) fm ON TRUE

  -- ... similar LATERAL joins for all other metric tables ...

  WHERE s.is_active = TRUE;""")

add_para('Why LATERAL? Each LATERAL join executes a subquery per row of market.securities, efficiently fetching only the most recent row from each time-series table without expensive correlated subqueries or separate CTEs.')

doc.add_page_break()

# ── STAGE 5 ──────────────────────────────────────────────────────────────────
add_heading('Stage 5 — API & Application Layer', level=2)

add_heading('Query Flow: Screener API', level=3)
add_code_block("""  GET /api/v1/screener?pe_max=20&sector=Banks&sort=dividend_yield&dir=desc

    →  SELECT * FROM screener.universe
       WHERE pe_ratio <= 20
         AND sector = 'Banks'
         AND is_active = TRUE
       ORDER BY dividend_yield DESC
       LIMIT 50 OFFSET 0

    →  Returns paginated JSON with all 458 columns (filtered server-side)
    →  No joins required — single-table scan with WHERE / ORDER BY / LIMIT""")

add_heading('Query Flow: Stock Detail API', level=3)
add_code_block("""  GET /api/v1/stocks/CBA

    →  screener.universe row           (quick summary)
    →  market.daily_prices             WHERE ticker='CBA' ORDER BY date DESC LIMIT 365
    →  market.fundamental_metrics      WHERE ticker='CBA' (all periods)
    →  market.dividends                WHERE ticker='CBA' ORDER BY ex_date DESC LIMIT 20
    →  market.announcements            WHERE ticker='CBA' ORDER BY released_at DESC LIMIT 10
    →  market.anomaly_flags            WHERE ticker='CBA' AND is_active=TRUE
    →  market.short_positions          WHERE ticker='CBA' ORDER BY report_date DESC LIMIT 1""")

doc.add_page_break()

# ── COMPLETE DATA LINEAGE MAP ─────────────────────────────────────────────────
add_heading('Complete Data Lineage Map', level=2)
add_code_block("""  EODHD API (JSON/CSV)
    ├── /eod/bulk          →  staging_au.bulk_eod          →  market.daily_prices
    │                                                          │
    │                                                          ├──►  market.daily_metrics
    │                                                          │       (returns, SMA, EMA, RSI,
    │                                                          │        MACD, BB, ATR, volume)
    │                                                          ├──►  market.weekly_metrics
    │                                                          ├──►  market.monthly_metrics
    │                                                          ├──►  market.quarterly_metrics
    │                                                          ├──►  market.halfyearly_metrics
    │                                                          └──►  market.yearly_metrics
    │
    ├── /fundamentals/{t}  →  staging_au.fundamentals      →  market.fundamentals
    │                                                          │
    ├── /dividends/{t}     →  staging_au.dividends    ──────────►│
    │                                                          ▼
    │                                                     market.fundamental_metrics
    │                                                     (PE, PB, EV/EBITDA, margins,
    │                                                      ROE, DY, franking, FCF yield)
    │
    ├── /splits/{t}        →  staging_au.splits            →  market.splits
    ├── /earnings/{t}      →  staging_au.earnings          →  market.earnings
    ├── /exchange-symbols  →  staging_au.exchange_symbols  →  market.securities
    └── /macro-indicator   →  staging_au.macro_indicators  →  market.macro_indicators

  ASIC CSV Download
    └── short_positions    →  staging_au.short_positions   →  market.short_positions

  ASX Announcements API
    └── announcement feed  →  (no staging, direct)      →  market.announcements

  market.daily_metrics + market.daily_prices
    └──────────────────────────────────────────────────►  market.anomaly_flags
                                                           (7 flag types, hourly)

  ALL MARKET TABLES
    └──────────────────────────────────────────────────►  screener.universe
                                                           (458 cols, nightly rebuild)
                                                                 │
                                                                 ▼
                                                           FastAPI /screener
                                                           FastAPI /stocks/{t}
                                                                 │
                                                                 ▼
                                                           Next.js Frontend""")

doc.add_page_break()

# ── SCHEDULING SUMMARY ────────────────────────────────────────────────────────
add_heading('Scheduling Summary', level=2)
add_table(
    headers=['Job', 'Frequency', 'Script / Scheduler', 'Key Operations'],
    rows=[
        ['Full daily ingest',     'After ASX close (~5pm AEST)', 'PM2 → daily_pipeline.py',   'Fetch prices → stage → compute all metrics → rebuild universe'],
        ['Fundamental refresh',   'Weekly (Monday)',             'PM2 → weekly_pipeline.py',  'Fetch fundamentals/splits/earnings → compute fundamental metrics'],
        ['Historical backfill',   'Monthly (1st)',               'PM2 → monthly_pipeline.py', 'Fill gaps, recompute full history'],
        ['Anomaly detection',     'Every 1 hour',               'APScheduler',               'Evaluate all 7 flag conditions, expire old flags'],
        ['Alert checking',        'Every 1 minute',             'APScheduler',               'Compare latest prices to user price alert thresholds'],
        ['Announcement fetch',    'Every 2 hours',              'APScheduler',               'Pull ASX announcement feed'],
        ['Short position update', 'Daily',                      'APScheduler',               'Download ASIC CSV, upsert positions'],
        ['Portfolio valuation',   'Every 5 minutes',            'APScheduler',               'Mark-to-market open portfolio holdings'],
    ],
    col_widths=[1.5, 1.5, 1.8, 3.2],
)

# ── KEY DESIGN PRINCIPLES ─────────────────────────────────────────────────────
add_heading('Key Design Principles', level=2)

principles = [
    ('1. Staging = Temporary, Canonical = Durable',
     'Staging tables are scratch space — always truncated before re-use. Market tables are the append-only source of truth.'),
    ('2. UPSERT everywhere in canonical layer',
     'ON CONFLICT (ticker, date) DO UPDATE ensures idempotent pipeline runs. Re-running the pipeline never creates duplicates.'),
    ('3. LATERAL joins for time-series "latest row"',
     'Far more efficient than correlated subqueries when joining many time-series tables to a securities master.'),
    ('4. Golden Record = query performance',
     'The screener universe is rebuilt nightly rather than computed at query time. A screener query with 20 filters on 2,000 stocks becomes a single-table scan, not a 10-way join.'),
    ('5. TimescaleDB for price data',
     'daily_prices and daily_metrics are TimescaleDB hypertables chunked by month. Range queries like WHERE date > NOW() - INTERVAL \'1 year\' skip irrelevant chunks entirely.'),
    ('6. Savepoint safety in pipeline',
     'Each _scalar() helper in the compute engine wraps its query in a PostgreSQL SAVEPOINT (BEGIN NESTED). A failed individual query rolls back only its own savepoint — it cannot poison the surrounding transaction and cause all subsequent metrics to silently return NULL/0.'),
    ('7. Trading-day lookbacks, not calendar days',
     'All return calculations use fixed trading-day offsets (1, 5, 21, 63, 126, 252) rather than calendar intervals. This ensures return_1m is always exactly 21 trading sessions regardless of public holidays or weekends.'),
]

for title, body in principles:
    p = doc.add_paragraph()
    run_bold = p.add_run(title)
    run_bold.font.bold = True
    run_bold.font.size = Pt(10.5)
    run_bold.font.name = 'Calibri'
    run_bold.font.color.rgb = RGBColor(0, 62, 107)
    run_body = p.add_run('\n' + body)
    run_body.font.size = Pt(10)
    run_body.font.name = 'Calibri'
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.space_after  = Pt(8)

# ── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f'Saved: {OUTPUT_PATH}')
