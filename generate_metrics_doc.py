"""
Generate the ASX Screener — Complete Metrics Reference Word document.
Organised as: Base Downloaded → Derived/Computed → Technical Calculated
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\Users\Dell\My Claude\Code\ASX Screener\ASX_Screener_Complete_Metrics_Reference.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)

# ── Helpers ───────────────────────────────────────────────────────────────────
BLUE_DARK  = (0, 62, 107)
BLUE_MID   = (0, 102, 160)
GREEN_DARK = (0, 100, 60)
PURPLE     = (80, 40, 120)
ORANGE     = (180, 80, 0)
RED_DARK   = (160, 20, 20)
TEAL       = (0, 100, 110)
GREY       = (80, 80, 80)

def rgb(t): return RGBColor(*t)

def add_heading(text, level=1, color=BLUE_DARK):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = rgb(color)
        run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_para(text='', bold=False, size=10, color=None, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color: run.font.color.rgb = rgb(color)
    return p

def add_separator(light=False):
    p = doc.add_paragraph('─' * 130)
    for run in p.runs:
        run.font.name  = 'Courier New'
        run.font.size  = Pt(6)
        run.font.color.rgb = rgb((200,200,200) if light else (160,160,160))
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

def tbl_shd(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None, hdr_fill='003E6B', stripe=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(8.5)
        run.font.name  = 'Calibri'
        run.font.color.rgb = rgb((255,255,255))
        tbl_shd(cell, hdr_fill)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        fill  = 'F0F4FA' if (stripe and ri % 2 == 0) else 'FFFFFF'
        for i, val in enumerate(row):
            cells[i].text = str(val)
            run = cells[i].paragraphs[0].runs[0] if cells[i].paragraphs[0].runs else cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.name = 'Calibri'
            tbl_shd(cells[i], fill)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ASX SCREENER')
run.font.name = 'Calibri'; run.font.size = Pt(30); run.font.bold = True
run.font.color.rgb = rgb(BLUE_DARK)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Complete Metrics Reference')
run2.font.name = 'Calibri'; run2.font.size = Pt(18)
run2.font.color.rgb = rgb(BLUE_MID)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Base Downloaded  ·  Derived  ·  Calculated\nFundamental  ·  Technical  ·  Composite Scores  ·  Anomaly Flags')
run3.font.name = 'Calibri'; run3.font.size = Pt(11)
run3.font.color.rgb = rgb(GREY)

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('Total: 500+ Metrics across 15 tables')
run4.font.name = 'Calibri'; run4.font.size = Pt(11); run4.font.bold = True
run4.font.color.rgb = rgb(GREEN_DARK)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Document Structure', level=1)
sections_overview = [
    ('SECTION 1', 'Base Metrics — Downloaded from EODHD / ASIC / ASX',   'Raw data exactly as received from external sources.'),
    ('SECTION 2', 'Daily Price & Technical Metrics',                       'market.daily_metrics — 90+ columns per trading day per stock.'),
    ('SECTION 3', 'Weekly Technical Metrics',                              'market.weekly_metrics — 45+ columns per week.'),
    ('SECTION 4', 'Monthly Technical Metrics',                             'market.monthly_metrics — 30+ columns per month.'),
    ('SECTION 5', 'Quarterly Fundamental Metrics',                         'market.quarterly_metrics — QoQ & YoY P&L comparisons.'),
    ('SECTION 6', 'Half-Yearly Fundamental Metrics',                       'market.halfyearly_metrics — HoH & YoY P&L comparisons.'),
    ('SECTION 7', 'Annual Fundamental Metrics (Ratios)',                   'market.yearly_metrics — 150+ computed fundamental ratios per year.'),
    ('SECTION 8', 'Raw Financial Statements',                              'financials.annual_pnl / balance_sheet / cashflow — source P&L, BS, CF.'),
    ('SECTION 9', 'Period Price Range Metrics',                            'market.period_metrics — High/Low/AvgVol across 7 lookback windows.'),
    ('SECTION 10','Composite Scores',                                      'screener.universe — Value, Quality, Growth, Momentum, Income, Composite.'),
    ('SECTION 11','Anomaly Flags',                                         'market.anomalies — 7 signal types with severity & lifecycle.'),
    ('SECTION 12','Screener Universe — Golden Record',                     'screener.universe — 458+ columns, nightly rebuild, all sections joined.'),
]
add_table(
    headers=['Section', 'Title', 'Description'],
    rows=sections_overview,
    col_widths=[0.8, 2.8, 4.0],
    hdr_fill='003E6B',
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — BASE METRICS DOWNLOADED
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 1 — Base Metrics Downloaded from External Sources', level=1, color=RED_DARK)
add_separator()
add_para('These are the raw fields ingested from EODHD API, ASIC, and ASX. No computation is applied. They land first in staging_au.* tables (TRUNCATE+RELOAD), then are upserted into canonical market.* and financials.* tables.', italic=True)

# 1.1 Daily Prices
add_heading('1.1  Daily Price Data  (market.daily_prices — TimescaleDB)', level=2, color=BLUE_MID)
add_para('Source: EODHD /eod/bulk API  |  Frequency: Daily after ASX close  |  Table: market.daily_prices')
add_table(
    headers=['Column', 'Type', 'Source Field', 'Description'],
    rows=[
        ['time',           'TIMESTAMPTZ', 'date',          'Trading date (TimescaleDB partition key)'],
        ['asx_code',       'TEXT',        'code',          'ASX ticker symbol (e.g. CBA, BHP)'],
        ['open',           'NUMERIC',     'open',          'Opening price for the session'],
        ['high',           'NUMERIC',     'high',          'Intraday high price'],
        ['low',            'NUMERIC',     'low',           'Intraday low price'],
        ['close',          'NUMERIC',     'close',         'Closing price'],
        ['adjusted_close', 'NUMERIC',     'adjusted_close','Dividend and split adjusted close'],
        ['volume',         'BIGINT',      'volume',        'Number of shares traded'],
        ['value_traded',   'NUMERIC',     'value',         'Dollar value of shares traded (price × volume)'],
        ['trades_count',   'INTEGER',     'trades',        'Number of discrete trades executed'],
        ['vwap',           'NUMERIC',     'vwap',          'Volume Weighted Average Price for the session'],
        ['data_source',    'TEXT',        '—',             'Source tag (eodhd_bulk / eodhd_eod)'],
    ],
    col_widths=[1.4, 1.0, 1.4, 3.8],
    hdr_fill='8B0000',
)

# 1.2 Annual P&L
add_heading('1.2  Annual Income Statement  (financials.annual_pnl)', level=2, color=BLUE_MID)
add_para('Source: EODHD /fundamentals — Highlights + Financials.Income_Statement  |  Frequency: Weekly  |  Schema: financials')
add_table(
    headers=['Column', 'Description'],
    rows=[
        ['asx_code',            'ASX ticker'],
        ['fiscal_year',         'Fiscal year (integer)'],
        ['period_end_date',     'Balance date for the period'],
        ['report_date',         'Date filed / released to ASX'],
        ['revenue',             'Total revenue (sales + other operating)'],
        ['cost_of_sales',       'Direct cost of goods/services sold'],
        ['gross_profit',        'Revenue minus cost of sales'],
        ['other_income',        'Non-operating / other income'],
        ['operating_expenses',  'SG&A, distribution, admin expenses'],
        ['ebitda',              'Earnings before interest, tax, depreciation, amortisation'],
        ['depreciation',        'Depreciation charge for the period'],
        ['amortisation',        'Amortisation of intangibles'],
        ['ebit',                'Earnings before interest and tax'],
        ['interest_expense',    'Finance costs / interest paid'],
        ['interest_income',     'Interest earned on cash/investments'],
        ['pbt',                 'Profit before tax'],
        ['tax',                 'Income tax expense'],
        ['pat',                 'Profit after tax (before minority)'],
        ['minority_interest',   'Earnings attributable to non-controlling interests'],
        ['net_profit',          'Net profit attributable to shareholders'],
        ['extraordinary_items', 'One-off / exceptional items'],
        ['dividend_paid',       'Total dividends paid during period'],
        ['material_cost',       'Raw material / input costs'],
        ['employee_cost',       'Staff wages and benefits'],
        ['eps',                 'Basic earnings per share'],
        ['eps_diluted',         'Diluted earnings per share'],
        ['shares_used',         'Weighted average shares used for EPS'],
        ['dps',                 'Dividends per share (declared/paid)'],
        ['dps_franking_pct',    'Franking percentage of dividend (0–100)'],
        ['dps_grossed_up',      'Dividend grossed up for franking credits'],
        ['opm',                 'Operating profit margin (EBIT/Revenue)'],
        ['npm',                 'Net profit margin (Net Profit/Revenue)'],
        ['gpm',                 'Gross profit margin (Gross Profit/Revenue)'],
        ['ebitda_margin',       'EBITDA as % of revenue'],
        ['currency',            'Reporting currency (AUD/USD/other)'],
        ['is_restated',         'Flag — restated/amended filing'],
        ['data_source',         'Source tag'],
    ],
    col_widths=[2.0, 5.6],
)

add_heading('1.3  Annual Balance Sheet  (financials.annual_balance_sheet)', level=2, color=BLUE_MID)
add_para('Source: EODHD /fundamentals — Balance_Sheet  |  Frequency: Weekly')
add_table(
    headers=['Column', 'Description'],
    rows=[
        ['cash_equivalents',       'Cash and short-term liquid assets'],
        ['trade_receivables',      'Accounts receivable from customers'],
        ['inventory',              'Stock on hand / work in progress'],
        ['other_current_assets',   'Prepayments, other short-term assets'],
        ['total_current_assets',   'Sum of all current assets'],
        ['gross_block',            'Property, plant & equipment (gross)'],
        ['accumulated_depreciation','Cumulative depreciation on fixed assets'],
        ['net_block',              'PP&E net of accumulated depreciation'],
        ['cwip',                   'Capital work in progress'],
        ['goodwill',               'Goodwill from acquisitions'],
        ['intangibles',            'Other intangible assets'],
        ['investments',            'Long-term investments / associates'],
        ['other_non_current',      'Other non-current assets'],
        ['total_assets',           'Total assets (current + non-current)'],
        ['trade_payables',         'Accounts payable to suppliers'],
        ['advance_from_customers', 'Deposits / advance payments received'],
        ['short_term_debt',        'Bank overdraft + current portion of LT debt'],
        ['other_current_liab',     'Accruals, other short-term liabilities'],
        ['total_current_liab',     'Sum of all current liabilities'],
        ['long_term_debt',         'Bonds, term loans, finance leases (non-current)'],
        ['lease_liabilities',      'AASB 16 / IFRS 16 lease obligations'],
        ['contingent_liabilities', 'Off-balance sheet contingent items'],
        ['other_non_current_liab', 'Deferred tax, other long-term liabilities'],
        ['total_liabilities',      'Total liabilities (current + non-current)'],
        ['equity_capital',         'Issued share capital (paid-up)'],
        ['preference_capital',     'Preference share capital'],
        ['reserves',               'General reserves and share premium'],
        ['retained_earnings',      'Accumulated retained profits'],
        ['minority_interest_bs',   'Non-controlling interest (balance sheet)'],
        ['total_equity',           'Total shareholders equity'],
        ['total_debt',             'Short-term debt + long-term debt'],
        ['net_debt',               'Total debt minus cash and equivalents'],
        ['working_capital',        'Current assets minus current liabilities'],
        ['book_value_per_share',   'Total equity / shares outstanding'],
        ['face_value',             'Par value / face value per share'],
        ['shares_outstanding',     'Total shares on issue'],
    ],
    col_widths=[2.0, 5.6],
)

add_heading('1.4  Annual Cash Flow Statement  (financials.annual_cashflow)', level=2, color=BLUE_MID)
add_para('Source: EODHD /fundamentals — Cash_Flow  |  Frequency: Weekly')
add_table(
    headers=['Column', 'Description'],
    rows=[
        ['net_income',            'Net profit (starting point for indirect method)'],
        ['depreciation_amort',    'Add-back of D&A (non-cash)'],
        ['working_capital_changes','Change in working capital items'],
        ['other_operating',       'Other operating adjustments'],
        ['cfo',                   'Cash Flow from Operations'],
        ['capex',                 'Capital expenditure (purchase of PP&E)'],
        ['acquisitions',          'Cash paid for business acquisitions'],
        ['disposals',             'Cash received from asset disposals'],
        ['investment_purchases',  'Purchase of financial investments'],
        ['other_investing',       'Other investing cash flows'],
        ['cfi',                   'Cash Flow from Investing'],
        ['dividends_paid',        'Dividends paid to shareholders'],
        ['debt_raised',           'Proceeds from new borrowings'],
        ['debt_repaid',           'Repayment of debt'],
        ['equity_raised',         'Proceeds from share issuance'],
        ['buybacks',              'Cash spent on share buybacks'],
        ['other_financing',       'Other financing cash flows'],
        ['cff',                   'Cash Flow from Financing'],
        ['net_change_in_cash',    'CFO + CFI + CFF'],
        ['opening_cash',          'Cash balance at start of period'],
        ['closing_cash',          'Cash balance at end of period'],
        ['fcf',                   'Free Cash Flow = CFO - CapEx'],
    ],
    col_widths=[2.0, 5.6],
)

add_heading('1.5  Half-Year P&L  (financials.half_year_pnl)', level=2, color=BLUE_MID)
add_para('Source: EODHD /fundamentals — half-year reporting  |  Key fields: period_end_date, period_label (H1/H2), revenue, gross_profit, other_income, ebitda, depreciation, ebit, interest_expense, pbt, tax, pat, net_profit, eps, dps, dps_franking_pct, opm, npm, gpm')

add_heading('1.6  Dividend Data  (market.dividends)', level=2, color=BLUE_MID)
add_table(
    headers=['Column', 'Description'],
    rows=[
        ['asx_code',       'ASX ticker'],
        ['ex_date',        'Ex-dividend date (ownership cut-off)'],
        ['payment_date',   'Actual cash payment date'],
        ['declaration_date','Date board declared the dividend'],
        ['amount',         'Dividend amount in AUD per share'],
        ['currency',       'Payment currency'],
        ['franking_pct',   'Franking percentage (0–100%)'],
        ['payment_type',   'cash / stock / special'],
        ['period',         'interim / final / special'],
        ['data_source',    'eodhd / asx_direct'],
    ],
    col_widths=[1.6, 6.0],
)

add_heading('1.7  Short Positions  (market.short_positions)', level=2, color=BLUE_MID)
add_para('Source: ASIC Daily Short Sale Report (CSV download)  |  Frequency: Daily')
add_table(
    headers=['Column', 'Description'],
    rows=[
        ['asx_code',             'ASX ticker'],
        ['report_date',          'ASIC reporting date'],
        ['short_position_shares','Total shares reported short'],
        ['total_shares',         'Total issued shares (for calculating %)'],
        ['short_position_pct',   'short_position_shares / total_shares × 100'],
    ],
    col_widths=[1.8, 5.8],
)

add_heading('1.8  ASX Announcements  (market.announcements)', level=2, color=BLUE_MID)
add_para('Source: ASX Announcements API (append-only, every 2 hours via APScheduler)')
add_table(
    headers=['Column', 'Description'],
    rows=[
        ['announcement_id',   'Unique ASX announcement identifier'],
        ['asx_code',          'ASX ticker'],
        ['released_at',       'Official release timestamp (ASX server time)'],
        ['headline',          'Announcement title/subject'],
        ['category',          'ASX category (Appendix 4C, Annual Report, etc.)'],
        ['is_price_sensitive','Boolean — flagged as price sensitive by ASX'],
        ['document_url',      'PDF document URL'],
        ['document_size',     'File size in bytes'],
    ],
    col_widths=[1.8, 5.8],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — DAILY TECHNICAL METRICS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 2 — Daily Price & Technical Metrics  (market.daily_metrics)', level=1, color=TEAL)
add_separator()
add_para('Engine: compute/engine/daily_metrics.py  |  Source: market.daily_prices  |  Frequency: Daily  |  Table type: TimescaleDB hypertable (partitioned by month)', italic=True)
add_para('Primary key: (asx_code, date)')

# 2.1 Price & Volume Base
add_heading('2.1  Price & Volume Fields', level=2, color=TEAL)
add_table(
    headers=['Column', 'Formula / Derivation', 'Description'],
    rows=[
        ['close',         'Direct from daily_prices',            'Unadjusted closing price'],
        ['adj_close',     'Direct from daily_prices',            'Dividend + split adjusted close'],
        ['open',          'Direct from daily_prices',            'Opening price (added migration 010)'],
        ['high',          'Direct from daily_prices',            'Intraday high (added migration 010)'],
        ['low',           'Direct from daily_prices',            'Intraday low (added migration 010)'],
        ['volume',        'Direct from daily_prices',            'Shares traded'],
        ['market_cap',    'close × shares_outstanding',          'Market capitalisation at close'],
        ['daily_return',  '(close - prev_close) / prev_close',   'Simple 1-day return'],
        ['log_return',    'LN(close / prev_close)',               'Log return (for volatility calc)'],
        ['gap_pct',       '(open - prev_close) / prev_close',    'Overnight gap % at open'],
        ['true_range',    'MAX(high-low, |high-prev_close|, |low-prev_close|)', 'True trading range'],
    ],
    col_widths=[1.5, 2.8, 3.3],
    hdr_fill='006B6B',
)

# 2.2 Moving Averages
add_heading('2.2  Moving Averages', level=2, color=TEAL)
add_table(
    headers=['Column', 'Window', 'Type', 'Formula'],
    rows=[
        ['sma_5',    '5d',   'Simple',      'AVG(close) over 5 days'],
        ['sma_10',   '10d',  'Simple',      'AVG(close) over 10 days'],
        ['sma_20',   '20d',  'Simple',      'AVG(close) over 20 days'],
        ['sma_50',   '50d',  'Simple',      'AVG(close) over 50 days'],
        ['sma_100',  '100d', 'Simple',      'AVG(close) over 100 days'],
        ['sma_200',  '200d', 'Simple',      'AVG(close) over 200 days'],
        ['ema_9',    '9d',   'Exponential', 'close × (2/10) + ema_9_prev × (8/10)'],
        ['ema_12',   '12d',  'Exponential', 'close × (2/13) + ema_12_prev × (11/13)'],
        ['ema_20',   '20d',  'Exponential', 'close × (2/21) + ema_20_prev × (19/21)'],
        ['ema_26',   '26d',  'Exponential', 'close × (2/27) + ema_26_prev × (25/27)'],
        ['ema_50',   '50d',  'Exponential', 'close × (2/51) + ema_50_prev × (49/51)'],
        ['ema_200',  '200d', 'Exponential', 'close × (2/201) + ema_200_prev × (199/201)'],
        ['sma_5_prev',  '5d-1',   'Simple', 'Prior day sma_5  (for crossover detection)'],
        ['sma_20_prev', '20d-1',  'Simple', 'Prior day sma_20 (for crossover detection)'],
        ['sma_50_prev', '50d-1',  'Simple', 'Prior day sma_50 (for golden/death cross)'],
        ['sma_200_prev','200d-1', 'Simple', 'Prior day sma_200 (for golden/death cross)'],
    ],
    col_widths=[1.2, 0.7, 1.2, 4.5],
    hdr_fill='006B6B',
)

# 2.3 MA Ratios & Signals
add_heading('2.3  Moving Average Ratios & Cross Signals', level=2, color=TEAL)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['dma20_ratio',    'close / sma_20 - 1',                    '% deviation from 20d SMA'],
        ['dma50_ratio',    'close / sma_50 - 1',                    '% deviation from 50d SMA'],
        ['dma200_ratio',   'close / sma_200 - 1',                   '% deviation from 200d SMA'],
        ['above_sma20',    '1 if close > sma_20 else 0',            'Binary: price above 20d SMA'],
        ['above_sma50',    '1 if close > sma_50 else 0',            'Binary: price above 50d SMA'],
        ['above_sma100',   '1 if close > sma_100 else 0',           'Binary: price above 100d SMA'],
        ['above_sma200',   '1 if close > sma_200 else 0',           'Binary: price above 200d SMA'],
        ['golden_cross',   '1 if sma_50 > sma_200 else 0',          'Bullish long-term trend (50d above 200d)'],
        ['death_cross',    '1 if sma_50 < sma_200 else 0',          'Bearish long-term trend (50d below 200d)'],
        ['macd_bullish_cross','1 if macd_line crosses above signal','MACD bullish signal (migration 010)'],
        ['macd_bearish_cross','1 if macd_line crosses below signal','MACD bearish signal (migration 010)'],
    ],
    col_widths=[1.6, 2.5, 3.5],
    hdr_fill='006B6B',
)

# 2.4 MACD
add_heading('2.4  MACD (Moving Average Convergence Divergence)', level=2, color=TEAL)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['macd_line',       'ema_12 - ema_26',                      'MACD line'],
        ['macd_signal',     'EMA(9) of macd_line',                  'Signal line (9-day EMA of MACD)'],
        ['macd_hist',       'macd_line - macd_signal',              'Histogram (momentum of momentum)'],
        ['macd_line_prev',  'macd_line[t-1]',                       'Prior day MACD line (crossover detection)'],
        ['macd_signal_prev','macd_signal[t-1]',                     'Prior day signal line'],
    ],
    col_widths=[1.6, 2.5, 3.5],
    hdr_fill='006B6B',
)

# 2.5 RSI
add_heading('2.5  RSI (Relative Strength Index)', level=2, color=TEAL)
add_table(
    headers=['Column', 'Period', 'Formula', 'Interpretation'],
    rows=[
        ['rsi_7',          '7d',  'RS = avg_gain_7d / avg_loss_7d ; RSI = 100 - 100/(1+RS)', 'Short-term momentum. >70 overbought, <30 oversold'],
        ['rsi_14',         '14d', 'Same with 14-day gains/losses',                           'Standard RSI. >70 overbought, <30 oversold'],
        ['rsi_21',         '21d', 'Same with 21-day gains/losses',                           'Smoother RSI. Less whipsaw'],
        ['rsi_overbought', '14d', '1 if rsi_14 > 70 else 0',                                 'Binary overbought flag (migration 010)'],
        ['rsi_oversold',   '14d', '1 if rsi_14 < 30 else 0',                                 'Binary oversold flag (migration 010)'],
    ],
    col_widths=[1.4, 0.6, 3.0, 2.6],
    hdr_fill='006B6B',
)

# 2.6 Bollinger Bands
add_heading('2.6  Bollinger Bands (20-day, 2σ)', level=2, color=TEAL)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['bb_mid',   'sma_20',                               'Middle band (20d SMA)'],
        ['bb_upper', 'sma_20 + 2 × STDDEV(close, 20d)',     'Upper band (2 standard deviations above)'],
        ['bb_lower', 'sma_20 - 2 × STDDEV(close, 20d)',     'Lower band (2 standard deviations below)'],
        ['bb_width', '(bb_upper - bb_lower) / bb_mid',      'Band width — normalised volatility measure'],
        ['bb_pct',   '(close - bb_lower) / (bb_upper - bb_lower)', 'Position within bands: 0=lower, 1=upper, >1=breakout above, <0=breakout below'],
    ],
    col_widths=[1.1, 2.8, 3.7],
    hdr_fill='006B6B',
)

# 2.7 Stochastics
add_heading('2.7  Stochastic Oscillator (14,3)', level=2, color=TEAL)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['stoch_k', '(close - low_14d) / (high_14d - low_14d) × 100', 'Fast %K: position of close in 14-day range (0–100)'],
        ['stoch_d', 'SMA(3) of stoch_k',                               'Slow %D: 3-day MA of %K (signal line)'],
    ],
    col_widths=[1.1, 3.5, 3.0],
    hdr_fill='006B6B',
)

# 2.8 Other Momentum/Trend
add_heading('2.8  Additional Momentum & Trend Indicators', level=2, color=TEAL)
add_table(
    headers=['Column', 'Formula / Window', 'Description'],
    rows=[
        ['adx_14',       'Average Directional Index (14d)',       'Trend strength 0–100. >25 = strong trend'],
        ['plus_di',      '+DI (14d)',                             'Positive directional indicator'],
        ['minus_di',     '-DI (14d)',                             'Negative directional indicator'],
        ['cci_20',       'CCI = (typical_price - sma_tp) / (0.015 × mean_dev_20d)', 'Commodity Channel Index. >100 overbought, <-100 oversold'],
        ['williams_r',   '(high_14d - close) / (high_14d - low_14d) × -100', 'Williams %R. 0 to -100. Above -20 overbought'],
        ['aroon_up',     '((14 - days_since_14d_high) / 14) × 100','Aroon Up — how recently 14d high was set'],
        ['aroon_down',   '((14 - days_since_14d_low) / 14) × 100', 'Aroon Down — how recently 14d low was set'],
        ['roc_10',       '(close - close[t-10]) / close[t-10] × 100', 'Rate of Change 10 days'],
        ['roc_20',       '(close - close[t-20]) / close[t-20] × 100', 'Rate of Change 20 days'],
        ['atr_14',       'EMA(14) of True Range',                 'Average True Range — volatility in price units'],
        ['atr_pct',      'atr_14 / close',                       'ATR as % of price — normalised volatility'],
    ],
    col_widths=[1.2, 2.8, 3.6],
    hdr_fill='006B6B',
)

# 2.9 Volume Indicators
add_heading('2.9  Volume Indicators', level=2, color=TEAL)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['volume_avg_5d',   'AVG(volume, 5d)',                   '5-day average volume'],
        ['volume_avg_20d',  'AVG(volume, 20d)',                  '20-day average volume'],
        ['volume_avg_50d',  'AVG(volume, 50d)',                  '50-day average volume'],
        ['volume_avg_52w',  'AVG(volume, 252d)',                 '52-week average daily volume (migration 010)'],
        ['relative_volume', 'volume / volume_avg_20d',          'Current volume vs 20d average (1.0 = normal)'],
        ['obv',             'OBV = cumulative (volume if up, -volume if down)', 'On-Balance Volume — trend confirmation via volume'],
        ['obv_ema',         'EMA(20) of OBV',                   'Smoothed OBV for signal generation'],
        ['vwap',            'SUM(price×volume) / SUM(volume) over session', 'Volume Weighted Average Price (intraday)'],
        ['cmf_20',          'Chaikin Money Flow 20d = SUM(MFV, 20d) / SUM(vol, 20d)', 'Money flow strength. +ve = accumulation, -ve = distribution'],
        ['mfi_14',          'Money Flow Index 14d (volume-weighted RSI)', 'Volume-weighted RSI. >80 overbought, <20 oversold'],
    ],
    col_widths=[1.5, 2.5, 3.6],
    hdr_fill='006B6B',
)

# 2.10 Volatility
add_heading('2.10  Volatility Metrics', level=2, color=TEAL)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['hv_20d',  'STDDEV(log_return, 20d) × SQRT(252)',  'Historical volatility annualised — 20-day window'],
        ['hv_60d',  'STDDEV(log_return, 60d) × SQRT(252)',  'Historical volatility annualised — 60-day window'],
    ],
    col_widths=[1.2, 3.0, 3.4],
    hdr_fill='006B6B',
)

# 2.11 Price Level & 52W
add_heading('2.11  Price Level & Reference Points', level=2, color=TEAL)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['high_52w',          'MAX(high, 252 trading days)',           '52-week high price'],
        ['low_52w',           'MIN(low, 252 trading days)',            '52-week low price'],
        ['ath_price',         'MAX(close, all history)',               'All-time high closing price'],
        ['atl_price',         'MIN(close, all history)',               'All-time low closing price'],
        ['pct_from_52w_high', '(close - high_52w) / high_52w',        '% below 52-week high (negative = down from high)'],
        ['pct_from_52w_low',  '(close - low_52w) / low_52w',          '% above 52-week low'],
        ['pct_from_ath',      '(close - ath_price) / ath_price',      '% below all-time high (drawdown from ATH)'],
        ['pct_from_atl',      '(close - atl_price) / atl_price',      '% above all-time low (migration 010)'],
        ['new_52w_high',      '1 if close = high_52w else 0',         'New 52-week high flag'],
        ['new_52w_low',       '1 if close = low_52w else 0',          'New 52-week low flag'],
        ['new_ath',           '1 if close = ath_price else 0',        'New all-time high flag'],
    ],
    col_widths=[1.7, 2.5, 3.4],
    hdr_fill='006B6B',
)

# 2.12 Return Lookbacks
add_heading('2.12  Return Lookbacks (Trading-Day Based)', level=2, color=TEAL)
add_para('Formula for all: (close[t] - close[t-N]) / close[t-N]  where N = trading days')
add_table(
    headers=['Column', 'N (trading days)', 'Calendar approx.'],
    rows=[
        ['return_1d',   '1',   '1 session'],
        ['return_5d',   '5',   '1 week  (migration 010)'],
        ['return_1w',   '5',   '1 week  (alias, migration 008)'],
        ['return_20d',  '20',  '1 month (migration 010)'],
        ['return_1m',   '21',  '1 month (migration 008)'],
        ['return_3m',   '63',  '3 months'],
        ['return_6m',   '126', '6 months'],
        ['return_ytd',  'From 1 Jan of current year', 'Year-to-date'],
        ['return_1y',   '252', '1 year'],
        ['return_60d',  '60',  '~3 months (migration 010)'],
    ],
    col_widths=[1.4, 1.5, 4.7],
    hdr_fill='006B6B',
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — WEEKLY METRICS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 3 — Weekly Technical Metrics  (market.weekly_metrics)', level=1, color=PURPLE)
add_separator()
add_para('Engine: compute/engine/weekly_metrics.py  |  Source: market.daily_prices  |  Primary key: (asx_code, week_date)', italic=True)

add_table(
    headers=['Column', 'Formula / Source', 'Description'],
    rows=[
        ['week_date',       'ISO Monday date',                          'Week start date (Monday)'],
        ['open',            'FIRST(open) of week',                      'Monday opening price'],
        ['high',            'MAX(high) of week',                        'Highest intraday price during week'],
        ['low',             'MIN(low) of week',                         'Lowest intraday price during week'],
        ['close',           'LAST(close) of week',                      'Friday closing price'],
        ['volume',          'SUM(volume) of week',                      'Total shares traded during week'],
        ['market_cap',      'close × shares_outstanding',               'Market cap at Friday close'],
        ['weekly_return',   '(close - prev_week_close) / prev_week_close','Week-over-week return'],
        ['return_4w',       '(close - close[t-4w]) / close[t-4w]',     '4-week (1-month) return'],
        ['return_13w',      '(close - close[t-13w]) / close[t-13w]',   '13-week (quarterly) return'],
        ['return_52w',      '(close - close[t-52w]) / close[t-52w]',   '52-week (annual) return'],
        ['adx_14',          'ADX(14) on weekly bars',                   'Trend strength on weekly timeframe'],
        ['plus_di',         '+DI(14) on weekly bars',                   'Positive directional movement'],
        ['minus_di',        '-DI(14) on weekly bars',                   'Negative directional movement'],
        ['aroon_up',        '((14 - weeks_since_14w_high) / 14) × 100','Recency of 14-week high'],
        ['aroon_down',      '((14 - weeks_since_14w_low) / 14) × 100', 'Recency of 14-week low'],
        ['rsi_14',          'RSI(14) on weekly closes',                 'Weekly RSI 14-period'],
        ['rsi_7',           'RSI(7) on weekly closes',                  'Weekly RSI 7-period (shorter)'],
        ['macd_line',       'EMA(13w) - EMA(26w)',                      'Weekly MACD line'],
        ['macd_signal',     'EMA(9) of weekly macd_line',               'Weekly MACD signal'],
        ['macd_hist',       'macd_line - macd_signal',                  'Weekly MACD histogram'],
        ['stoch_k',         '(close - low_14w) / (high_14w - low_14w) × 100','Weekly Stochastic %K'],
        ['stoch_d',         'SMA(3) of stoch_k',                        'Weekly Stochastic %D'],
        ['cci_20',          'CCI(20) on weekly bars',                   'Weekly CCI 20-period'],
        ['williams_r',      'Williams %R(14) on weekly bars',           'Weekly Williams %R'],
        ['sma_10w',         'AVG(close, 10 weeks)',                     '10-week SMA'],
        ['sma_20w',         'AVG(close, 20 weeks)',                     '20-week SMA (~5 months)'],
        ['sma_40w',         'AVG(close, 40 weeks)',                     '40-week SMA (~10 months)'],
        ['ema_13w',         'EMA(13) of weekly closes',                 '13-week EMA'],
        ['ema_26w',         'EMA(26) of weekly closes',                 '26-week EMA'],
        ['atr_14',          'ATR(14) on weekly bars',                   'Weekly Average True Range'],
        ['bb_upper',        'sma_20w + 2σ',                             'Weekly Bollinger upper band'],
        ['bb_lower',        'sma_20w - 2σ',                             'Weekly Bollinger lower band'],
        ['bb_pct',          '(close - bb_lower) / (bb_upper - bb_lower)','Position within weekly BB'],
        ['bb_width',        '(bb_upper - bb_lower) / sma_20w',         'Weekly BB normalised width'],
        ['volume_avg_4w',   'AVG(volume, 4 weeks)',                     '4-week average volume'],
        ['relative_volume', 'volume / volume_avg_4w',                   'Volume vs 4-week average'],
        ['obv',             'Cumulative OBV on weekly bars',            'Weekly On-Balance Volume'],
        ['golden_cross',    '1 if sma_10w > sma_40w',                  'Weekly golden cross (10w above 40w)'],
        ['death_cross',     '1 if sma_10w < sma_40w',                  'Weekly death cross'],
        ['above_sma10w',    '1 if close > sma_10w',                    'Price above 10-week SMA'],
        ['above_sma40w',    '1 if close > sma_40w',                    'Price above 40-week SMA'],
    ],
    col_widths=[1.5, 2.5, 3.6],
    hdr_fill='502878',
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — MONTHLY METRICS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 4 — Monthly Technical Metrics  (market.monthly_metrics)', level=1, color=ORANGE)
add_separator()
add_para('Engine: compute/engine/monthly_metrics.py  |  Source: market.daily_prices  |  Primary key: (asx_code, month_date)', italic=True)
add_table(
    headers=['Column', 'Formula / Source', 'Description'],
    rows=[
        ['month_date',          'First day of calendar month',                      'Month reference date'],
        ['close',               'LAST(close) of month',                             'Month-end closing price'],
        ['volume_avg',          'AVG(daily volume) over month',                     'Average daily volume for month'],
        ['market_cap',          'close × shares_outstanding',                       'Market cap at month end'],
        ['monthly_return',      '(close - prev_month_close) / prev_month_close',    'Month-over-month return'],
        ['return_3m',           '(close - close[t-3m]) / close[t-3m]',             '3-month trailing return'],
        ['return_6m',           '(close - close[t-6m]) / close[t-6m]',             '6-month trailing return'],
        ['return_12m',          '(close - close[t-12m]) / close[t-12m]',           '12-month (1-year) return'],
        ['return_ytd',          '(close - close[Jan 1]) / close[Jan 1]',            'Year-to-date return'],
        ['momentum_3m',         'return_3m (used in momentum factor scoring)',       '3-month momentum factor'],
        ['momentum_6m',         'return_6m',                                        '6-month momentum factor'],
        ['momentum_12m',        'return_12m',                                       '12-month momentum factor (Jegadeesh-Titman)'],
        ['relative_strength_xjo','return_12m - ASX200_return_12m',                 'Return vs ASX 200 index (alpha proxy)'],
        ['volatility_1m',       'STDDEV(daily_return, 21d) × SQRT(252)',            'Annualised vol from 1-month daily returns'],
        ['volatility_3m',       'STDDEV(daily_return, 63d) × SQRT(252)',            'Annualised vol from 3-month daily returns'],
        ['volatility_12m',      'STDDEV(daily_return, 252d) × SQRT(252)',           'Annualised vol from 1-year daily returns'],
        ['rsi_14',              'RSI(14) computed on monthly closes',               'Monthly RSI 14-period'],
        ['macd_line',           'EMA(12m) - EMA(26m)',                              'Monthly MACD line'],
        ['macd_signal',         'EMA(9) of monthly MACD',                           'Monthly MACD signal'],
        ['macd_hist',           'macd_line - macd_signal',                          'Monthly MACD histogram'],
        ['bb_pct',              '(close - bb_lower_20m) / (bb_upper_20m - bb_lower_20m)', 'Monthly BB position'],
        ['bb_width',            '(bb_upper - bb_lower) / sma_12m',                  'Monthly BB normalised width'],
        ['sma_12m',             'AVG(close, 12 months)',                             '12-month SMA'],
        ['price_to_52w_high',   'close / MAX(close, 252d)',                         'Price as fraction of 52-week high'],
        ['drawdown_from_ath',   '(close - ath_price) / ath_price',                  'Drawdown % from all-time high'],
    ],
    col_widths=[1.7, 2.8, 3.1],
    hdr_fill='B45000',
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — QUARTERLY FUNDAMENTAL METRICS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 5 — Quarterly Fundamental Metrics  (market.quarterly_metrics)', level=1, color=GREEN_DARK)
add_separator()
add_para('Source: financials.half_year_pnl + EODHD quarterly reports  |  Primary key: (asx_code, fiscal_year, quarter)', italic=True)
add_table(
    headers=['Column', 'Type', 'Description'],
    rows=[
        ['fiscal_year',          'Raw',    'Fiscal year of the period'],
        ['quarter',              'Raw',    'Quarter number (1–4)'],
        ['period_end_date',      'Raw',    'End date of the quarter'],
        ['period_label',         'Raw',    'Label e.g. Q1 FY2025'],
        ['revenue',              'Raw',    'Quarterly revenue'],
        ['gross_profit',         'Raw',    'Quarterly gross profit'],
        ['ebitda',               'Raw',    'Quarterly EBITDA'],
        ['ebit',                 'Raw',    'Quarterly EBIT'],
        ['other_income',         'Raw',    'Other income for quarter'],
        ['interest_expense',     'Raw',    'Interest expense for quarter'],
        ['depreciation',         'Raw',    'D&A for quarter'],
        ['tax',                  'Raw',    'Tax charge for quarter'],
        ['net_income',           'Raw',    'Net profit for quarter'],
        ['extraordinary_items',  'Raw',    'One-off items'],
        ['equity_capital',       'Raw',    'Shares on issue'],
        ['eps',                  'Raw',    'Quarterly EPS'],
        ['gross_margin',         'Derived','gross_profit / revenue'],
        ['ebit_margin',          'Derived','ebit / revenue'],
        ['net_margin',           'Derived','net_income / revenue'],
        ['revenue_growth_qoq',   'Derived','(revenue - prior_q_revenue) / prior_q_revenue'],
        ['net_income_growth_qoq','Derived','(net_income - prior_q_net_income) / prior_q_net_income'],
        ['ebit_growth_qoq',      'Derived','(ebit - prior_q_ebit) / prior_q_ebit'],
        ['revenue_growth_yoy',   'Derived','(revenue - same_q_prior_yr_revenue) / same_q_prior_yr_revenue'],
        ['net_income_growth_yoy','Derived','(net_income - same_q_prior_yr_net_income) / same_q_prior_yr_net_income'],
        ['ebit_growth_yoy',      'Derived','(ebit - same_q_prior_yr_ebit) / same_q_prior_yr_ebit'],
        ['eps_growth_yoy',       'Derived','(eps - same_q_prior_yr_eps) / same_q_prior_yr_eps'],
    ],
    col_widths=[2.0, 0.8, 4.8],
    hdr_fill='006432',
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — HALF-YEARLY METRICS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 6 — Half-Yearly Fundamental Metrics  (market.halfyearly_metrics)', level=1, color=GREEN_DARK)
add_separator()
add_para('Source: financials.half_year_pnl  |  Primary key: (asx_code, period_end_date)', italic=True)
add_table(
    headers=['Column', 'Type', 'Description'],
    rows=[
        ['period_end_date',      'Raw',    'Half-year end date'],
        ['fiscal_year',          'Raw',    'Fiscal year'],
        ['half',                 'Raw',    '1 = H1 (Jan-Jun), 2 = H2 (Jul-Dec)'],
        ['period_label',         'Raw',    'e.g. H1 FY2025'],
        ['revenue',              'Raw',    'Half-year revenue'],
        ['gross_profit',         'Raw',    'Half-year gross profit'],
        ['ebitda',               'Raw',    'Half-year EBITDA'],
        ['ebit',                 'Raw',    'Half-year EBIT'],
        ['net_income',           'Raw',    'Half-year net profit'],
        ['other_income',         'Raw',    'Other income'],
        ['interest_expense',     'Raw',    'Finance costs'],
        ['tax',                  'Raw',    'Tax charge'],
        ['depreciation',         'Raw',    'D&A'],
        ['eps',                  'Raw',    'Half-year EPS'],
        ['dps',                  'Raw',    'Half-year DPS'],
        ['franking_pct',         'Raw',    'Dividend franking %'],
        ['gross_margin',         'Derived','gross_profit / revenue'],
        ['ebit_margin',          'Derived','ebit / revenue'],
        ['net_margin',           'Derived','net_income / revenue'],
        ['revenue_growth_hoh',   'Derived','(revenue - prior_half_revenue) / prior_half_revenue (Half-on-Half)'],
        ['revenue_growth_yoy',   'Derived','(revenue - same_half_prior_yr) / same_half_prior_yr'],
        ['net_income_growth_hoh','Derived','Half-on-half net income growth'],
        ['net_income_growth_yoy','Derived','Same-half year-on-year net income growth'],
        ['eps_growth_hoh',       'Derived','Half-on-half EPS growth'],
        ['eps_growth_yoy',       'Derived','Same-half year-on-year EPS growth'],
    ],
    col_widths=[2.0, 0.8, 4.8],
    hdr_fill='006432',
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — ANNUAL FUNDAMENTAL RATIOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 7 — Annual Fundamental Ratios  (market.yearly_metrics)', level=1, color=BLUE_DARK)
add_separator()
add_para('Engine: compute/engine/fundamental_metrics.py  |  Sources: financials.* + market.daily_prices + market.dividends  |  Primary key: (asx_code, fiscal_year)', italic=True)

add_heading('7.1  Price & Market Reference', level=2, color=BLUE_MID)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['price_at_compute',  'Latest close at time of compute',     'Reference price used for ratio computation'],
        ['market_cap',        'close × shares_outstanding',          'Market capitalisation'],
        ['enterprise_value',  'market_cap + total_debt - cash',      'EV = market cap + net debt'],
        ['shares_outstanding','From balance sheet',                   'Total shares on issue'],
    ],
    col_widths=[1.8, 2.5, 3.3],
    hdr_fill='003E6B',
)

add_heading('7.2  Valuation Ratios', level=2, color=BLUE_MID)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['pe_ratio',        'price / eps_diluted',              'Price-to-Earnings (P/E)'],
        ['pb_ratio',        'price / bvps',                     'Price-to-Book (P/B)'],
        ['ps_ratio',        'market_cap / revenue',             'Price-to-Sales (P/S)'],
        ['pcf_ratio',       'market_cap / cfo',                 'Price-to-Cash Flow'],
        ['p_fcf_ratio',     'market_cap / fcf',                 'Price-to-Free Cash Flow'],
        ['ev_ebitda',       'enterprise_value / ebitda',        'EV/EBITDA'],
        ['ev_ebit',         'enterprise_value / ebit',          'EV/EBIT'],
        ['ev_revenue',      'enterprise_value / revenue',       'EV/Revenue'],
        ['ev_fcf',          'enterprise_value / fcf',           'EV/Free Cash Flow'],
        ['peg_ratio',       'pe_ratio / eps_growth_1y',         'Price/Earnings-to-Growth'],
        ['earnings_yield',  'eps_diluted / price',              'Inverse P/E (earnings yield %)'],
        ['fcf_yield',       'fcf / market_cap',                 'Free cash flow yield'],
        ['graham_number',   'SQRT(22.5 × eps × bvps)',          'Graham intrinsic value estimate'],
        ['ncavps',          '(current_assets - total_liab) / shares', 'Net Current Asset Value per Share'],
        ['pe_5y_avg',       'AVG(pe_ratio, last 5 years)',      '5-year average P/E'],
    ],
    col_widths=[1.5, 2.5, 3.6],
    hdr_fill='003E6B',
)

add_heading('7.3  Per Share Data', level=2, color=BLUE_MID)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['eps',             'net_profit / weighted_avg_shares',   'Basic EPS'],
        ['eps_diluted',     'net_profit / diluted_shares',        'Diluted EPS'],
        ['bvps',            'total_equity / shares_outstanding',  'Book Value per Share'],
        ['tbvps',           '(total_equity - intangibles - goodwill) / shares', 'Tangible Book Value per Share'],
        ['dps',             'dividends_paid / shares',            'Dividends per Share'],
        ['dps_grossed_up',  'dps + franking_credit_per_share',   'Grossed-up DPS (incl. franking credit)'],
        ['fcf_per_share',   'fcf / shares_outstanding',           'Free Cash Flow per Share'],
        ['ocf_per_share',   'cfo / shares_outstanding',           'Operating Cash Flow per Share'],
        ['revenue_per_share','revenue / shares_outstanding',      'Revenue per Share'],
        ['net_debt_per_share','net_debt / shares_outstanding',    'Net Debt per Share'],
    ],
    col_widths=[1.6, 2.5, 3.5],
    hdr_fill='003E6B',
)

add_heading('7.4  Dividend & Franking Metrics', level=2, color=BLUE_MID)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['dividend_yield',         'dps / price × 100',                                'Dividend yield %'],
        ['franking_pct',           'From market.dividends',                             'Franking % (0–100)'],
        ['franked_yield',          'dps_grossed_up / price × 100',                     'Grossed-up yield (incl. franking credits)'],
        ['payout_ratio',           'dps / eps × 100',                                  'Dividends as % of earnings'],
        ['dividend_cagr_3y',       'CAGR(dps, 3 years)',                               '3-year dividend per share CAGR'],
        ['dividend_cagr_5y',       'CAGR(dps, 5 years)',                               '5-year dividend per share CAGR'],
        ['dividend_consecutive_yrs','COUNT(consecutive years with div payment)',        'Years of consecutive dividend payments'],
    ],
    col_widths=[1.9, 2.5, 3.2],
    hdr_fill='003E6B',
)
add_para('Franking Credit per Share = dps × (franking_pct/100) × (0.30/0.70)', italic=True, size=9)

add_heading('7.5  Profitability & Returns', level=2, color=BLUE_MID)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['roe',            'net_profit / total_equity',              'Return on Equity'],
        ['roa',            'net_profit / total_assets',              'Return on Assets'],
        ['roic',           'nopat / invested_capital',               'Return on Invested Capital'],
        ['roce',           'ebit / (total_assets - current_liab)',   'Return on Capital Employed'],
        ['roae',           'net_profit / avg_equity',                'Return on Average Equity'],
        ['roaa',           'net_profit / avg_assets',                'Return on Average Assets'],
        ['croic',          'fcf / invested_capital',                 'Cash Return on Invested Capital'],
    ],
    col_widths=[1.2, 2.5, 3.9],
    hdr_fill='003E6B',
)

add_heading('7.6  Margin Ratios', level=2, color=BLUE_MID)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['gross_margin',     'gross_profit / revenue',    'Gross profit margin'],
        ['ebitda_margin',    'ebitda / revenue',          'EBITDA margin'],
        ['ebit_margin',      'ebit / revenue',            'EBIT (operating) margin'],
        ['pretax_margin',    'pbt / revenue',             'Pre-tax profit margin'],
        ['net_margin',       'net_profit / revenue',      'Net profit margin'],
        ['ocf_margin',       'cfo / revenue',             'Operating cash flow margin'],
        ['fcf_margin',       'fcf / revenue',             'Free cash flow margin'],
        ['tax_rate_effective','tax / pbt',                'Effective tax rate'],
    ],
    col_widths=[1.5, 2.0, 4.1],
    hdr_fill='003E6B',
)

add_heading('7.7  Efficiency Ratios', level=2, color=BLUE_MID)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['asset_turnover',      'revenue / total_assets',             'Asset utilisation efficiency'],
        ['inventory_turnover',  'cogs / inventory',                   'Inventory cycle speed'],
        ['receivables_turnover','revenue / trade_receivables',        'Receivables collection speed'],
        ['receivables_days',    '365 / receivables_turnover',         'Days Sales Outstanding (DSO)'],
        ['inventory_days',      '365 / inventory_turnover',           'Days Inventory Outstanding (DIO)'],
        ['payables_days',       '(trade_payables / cogs) × 365',      'Days Payables Outstanding (DPO)'],
        ['cash_conversion_cycle','receivables_days + inventory_days - payables_days', 'CCC = net days to convert investments to cash'],
        ['capex_intensity',     'capex / revenue',                    'CapEx as % of revenue'],
        ['revenue_per_employee','revenue / headcount',                'Revenue per employee (if available)'],
    ],
    col_widths=[1.8, 2.5, 3.3],
    hdr_fill='003E6B',
)

add_heading('7.8  Liquidity & Leverage Ratios', level=2, color=BLUE_MID)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['current_ratio',      'current_assets / current_liab',          'Ability to cover short-term obligations'],
        ['quick_ratio',        '(current_assets - inventory) / current_liab', 'Current ratio ex-inventory'],
        ['cash_ratio',         'cash / current_liab',                    'Most conservative liquidity measure'],
        ['debt_to_equity',     'total_debt / total_equity',              'Financial leverage ratio'],
        ['debt_to_assets',     'total_debt / total_assets',              'Asset financing from debt'],
        ['debt_to_ebitda',     'total_debt / ebitda',                    'Debt serviceability'],
        ['net_debt_to_ebitda', 'net_debt / ebitda',                      'Net debt coverage'],
        ['net_debt_to_equity', 'net_debt / total_equity',                'Gearing ratio'],
        ['interest_coverage',  'ebit / interest_expense',                'Times interest earned'],
        ['equity_multiplier',  'total_assets / total_equity',            'DuPont leverage component'],
        ['lt_debt_to_capital', 'long_term_debt / (long_term_debt + equity)', 'Long-term debt proportion of capital'],
    ],
    col_widths=[1.8, 2.5, 3.3],
    hdr_fill='003E6B',
)

add_heading('7.9  Quality Scores', level=2, color=BLUE_MID)
add_table(
    headers=['Column', 'Formula / Description', 'Range'],
    rows=[
        ['piotroski_f_score',  '9-point score: 3×profitability + 3×leverage/liquidity + 3×operating efficiency signals. Each sub-test = 0 or 1.', '0–9 (8-9 = strong)'],
        ['altman_z_score',     'Z = 1.2×X1 + 1.4×X2 + 3.3×X3 + 0.6×X4 + 1.0×X5 where X1–X5 are WC/TA, RE/TA, EBIT/TA, MV/TL, Sales/TA', '>2.99 safe, 1.81–2.99 grey, <1.81 distress'],
        ['beneish_m_score',    '8-variable model to detect earnings manipulation. M = -4.84 + weighted sum of 8 financial ratios.', '<-1.78 likely manipulator'],
    ],
    col_widths=[1.6, 4.3, 1.7],
    hdr_fill='003E6B',
)

add_heading('7.10  Growth Rates (Annual)', level=2, color=BLUE_MID)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['revenue_growth_1y',       '(rev_yr - rev_yr-1) / rev_yr-1',       'Revenue year-on-year growth'],
        ['gross_profit_growth_1y',  'GP growth YoY',                         'Gross profit YoY growth'],
        ['ebitda_growth_1y',        'EBITDA growth YoY',                     'EBITDA YoY growth'],
        ['ebit_growth_1y',          'EBIT growth YoY',                       'EBIT YoY growth'],
        ['net_income_growth_1y',    'NI growth YoY',                         'Net income YoY growth'],
        ['eps_growth_1y',           'EPS growth YoY',                        'EPS YoY growth'],
        ['ocf_growth_1y',           'CFO growth YoY',                        'Operating cashflow YoY growth'],
        ['fcf_growth_1y',           'FCF growth YoY',                        'Free cashflow YoY growth'],
        ['bvps_growth_1y',          'BVPS growth YoY',                       'Book value per share YoY growth'],
    ],
    col_widths=[2.0, 2.0, 3.6],
    hdr_fill='003E6B',
)

add_heading('7.11  Multi-Year CAGRs', level=2, color=BLUE_MID)
add_table(
    headers=['Metric', '3Y CAGR', '5Y CAGR', '7Y CAGR', '10Y CAGR'],
    rows=[
        ['Revenue',     'revenue_cagr_3y',    'revenue_cagr_5y',    'revenue_cagr_7y',    'revenue_cagr_10y'],
        ['Net Income',  'net_income_cagr_3y', 'net_income_cagr_5y', 'net_income_cagr_7y', 'net_income_cagr_10y'],
        ['EPS',         'eps_cagr_3y',        'eps_cagr_5y',        'eps_cagr_7y',        'eps_cagr_10y'],
        ['EBITDA',      'ebitda_cagr_3y',     'ebitda_cagr_5y',     'ebitda_cagr_7y',     'ebitda_cagr_10y'],
        ['FCF',         'fcf_cagr_3y',        'fcf_cagr_5y',        '—',                  '—'],
        ['Gross Profit','gross_profit_cagr_3y','gross_profit_cagr_5y','—',                 '—'],
        ['BVPS',        'bvps_cagr_3y',       'bvps_cagr_5y',       '—',                  '—'],
        ['Price',       'price_cagr_1y',      'price_cagr_3y / price_cagr_5y', 'price_cagr_7y', 'price_cagr_10y'],
        ['Rev (median)','—',                  'revenue_growth_median_5y','—',              'revenue_growth_median_10y'],
    ],
    col_widths=[1.3, 1.5, 1.5, 1.5, 1.8],
    hdr_fill='003E6B',
)
add_para('Formula: CAGR = (value_end / value_start)^(1/N) - 1')

add_heading('7.12  Multi-Year Averages', level=2, color=BLUE_MID)
add_table(
    headers=['Metric', '3Y Avg', '5Y Avg', '7Y Avg', '10Y Avg'],
    rows=[
        ['ROE',             'avg_roe_3y',          'avg_roe_5y',          'avg_roe_7y',          'avg_roe_10y'],
        ['ROA',             'avg_roa_3y',          'avg_roa_5y',          '—',                   '—'],
        ['ROIC',            'avg_roic_3y',         'avg_roic_5y',         '—',                   '—'],
        ['ROCE',            'avg_roce_3y',         'avg_roce_5y',         'avg_roce_7y',         'avg_roce_10y'],
        ['Gross Margin',    'avg_gross_margin_3y', 'avg_gross_margin_5y', '—',                   '—'],
        ['EBITDA Margin',   'avg_ebitda_margin_3y','avg_ebitda_margin_5y','—',                   '—'],
        ['Operating Margin','avg_operating_margin_3y','avg_operating_margin_5y','—',             'avg_operating_margin_10y'],
        ['Net Margin',      'avg_net_margin_3y',   'avg_net_margin_5y',   '—',                   'avg_net_margin_10y'],
        ['FCF Margin',      'avg_fcf_margin_3y',   'avg_fcf_margin_5y',   '—',                   '—'],
        ['EPS Growth',      'avg_eps_growth_3y',   'avg_eps_growth_5y',   '—',                   'avg_eps_growth_10y'],
        ['Current Ratio',   'avg_current_ratio_3y','—',                   '—',                   '—'],
        ['EBIT',            '—',                   'avg_ebit_5y',         '—',                   'avg_ebit_10y'],
    ],
    col_widths=[1.6, 1.5, 1.5, 1.3, 1.7],
    hdr_fill='003E6B',
)

add_heading('7.13  Risk & Price Return Metrics', level=2, color=BLUE_MID)
add_table(
    headers=['Column', 'Formula', 'Description'],
    rows=[
        ['beta_1y',       'Cov(stock, XJO_1y) / Var(XJO_1y)',      'Beta vs ASX 200 over 1 year'],
        ['beta_3y',       'Cov(stock, XJO_3y) / Var(XJO_3y)',      'Beta vs ASX 200 over 3 years'],
        ['beta_5y',       'Cov(stock, XJO_5y) / Var(XJO_5y)',      'Beta vs ASX 200 over 5 years'],
        ['volatility_1y', 'STDDEV(daily_return, 252d) × √252',     'Annualised 1-year volatility'],
        ['volatility_3y', 'STDDEV(daily_return, 756d) × √252',     'Annualised 3-year volatility'],
        ['sharpe_1y',     '(return_1y - risk_free_rate) / vol_1y', 'Sharpe ratio 1-year'],
        ['sharpe_3y',     '(return_3y - risk_free_rate) / vol_3y', 'Sharpe ratio 3-year'],
        ['sortino_1y',    '(return_1y - risk_free_rate) / downside_vol_1y', 'Sortino ratio (downside vol only)'],
        ['max_drawdown_1y','Max peak-to-trough decline (1 year)',  'Worst loss from peak in 1-year window'],
        ['max_drawdown_3y','Max peak-to-trough decline (3 years)', 'Worst loss from peak in 3-year window'],
        ['calmar_ratio',  'return_3y / max_drawdown_3y',           'Return per unit of max drawdown'],
        ['alpha_1y',      'return_1y - (beta_1y × XJO_return_1y)','1-year Jensen\'s alpha vs ASX 200'],
        ['alpha_3y',      'return_3y - (beta_3y × XJO_return_3y)','3-year Jensen\'s alpha vs ASX 200'],
        ['return_1y',     '(price_end - price_start) / price_start','1-year price return'],
        ['return_3y',     '3-year total return (migration 017)',    '3-year price return CAGR'],
        ['return_5y',     '5-year total return (migration 017)',    '5-year price return CAGR'],
        ['return_7y',     '7-year total return (migration 017)',    '7-year price return CAGR'],
        ['return_10y',    '10-year total return (migration 017)',   '10-year price return CAGR'],
        ['return_15y',    '15-year total return (migration 017)',   '15-year price return CAGR'],
    ],
    col_widths=[1.5, 2.5, 3.6],
    hdr_fill='003E6B',
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — PERIOD PRICE RANGE METRICS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 8 — Period Price Range Metrics  (market.period_metrics)', level=1, color=TEAL)
add_separator()
add_para('Source: market.daily_prices  |  Primary key: (asx_code, computed_date)  |  Rebuilt nightly', italic=True)
add_table(
    headers=['Column', 'Window', 'Description'],
    rows=[
        ['high_1d',       '1 trading day',  'Yesterday\'s high'],
        ['low_1d',        '1 trading day',  'Yesterday\'s low'],
        ['avg_volume_1d', '1 trading day',  'Yesterday\'s volume'],
        ['high_1w',       '5 trading days', '1-week price high'],
        ['low_1w',        '5 trading days', '1-week price low'],
        ['avg_volume_1w', '5 trading days', '1-week average daily volume'],
        ['high_1m',       '21 trading days','1-month price high'],
        ['low_1m',        '21 trading days','1-month price low'],
        ['avg_volume_1m', '21 trading days','1-month average daily volume'],
        ['high_3m',       '63 trading days','3-month price high'],
        ['low_3m',        '63 trading days','3-month price low'],
        ['avg_volume_3m', '63 trading days','3-month average daily volume'],
        ['high_6m',       '126 trading days','6-month price high'],
        ['low_6m',        '126 trading days','6-month price low'],
        ['avg_volume_6m', '126 trading days','6-month average daily volume'],
        ['high_1y',       '252 trading days','1-year price high'],
        ['low_1y',        '252 trading days','1-year price low'],
        ['avg_volume_1y', '252 trading days','1-year average daily volume'],
        ['high_52w',      '252 trading days','52-week high (alias of high_1y)'],
        ['low_52w',       '252 trading days','52-week low (alias of low_1y)'],
        ['avg_volume_52w','252 trading days','52-week average daily volume'],
    ],
    col_widths=[1.6, 1.4, 4.6],
    hdr_fill='006B6B',
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — COMPOSITE SCORES
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 9 — Composite Scores  (screener.universe)', level=1, color=PURPLE)
add_separator()
add_para('Added by migration 023_composite_scores.sql  |  Range: 0–100 (percentile rank within ASX universe)  |  Recomputed nightly', italic=True)
add_table(
    headers=['Score', 'Key Input Metrics', 'Description'],
    rows=[
        ['value_score',     'pe_ratio, pb_ratio, ps_ratio, ev_ebitda, fcf_yield, earnings_yield',                 'Composite value factor. Higher = cheaper relative to fundamentals.'],
        ['quality_score',   'roe, roic, gross_margin, net_margin, debt_to_equity, piotroski_f_score, altman_z',   'Composite quality factor. Higher = more profitable, lower leverage, better financial health.'],
        ['growth_score',    'revenue_growth_yoy, eps_growth_1y, revenue_cagr_3y, eps_cagr_3y, fcf_growth_1y',     'Composite growth factor. Higher = faster-growing company.'],
        ['momentum_score',  'return_1m, return_3m, return_6m, return_12m, rsi_14, relative_strength_xjo',         'Composite momentum factor. Higher = stronger recent price performance.'],
        ['income_score',    'dividend_yield, grossed_up_yield, payout_ratio, dividend_consecutive_yrs, fcf_yield','Composite income factor. Higher = better dividend & income characteristics.'],
        ['composite_score', 'Weighted average of all 5 factor scores',                                            'Overall composite score. Blends value, quality, growth, momentum and income.'],
    ],
    col_widths=[1.4, 3.2, 3.0],
    hdr_fill='502878',
)
add_para('All scores are percentile-ranked within the active ASX universe on each nightly rebuild. A score of 80 means the stock ranks in the top 20% of all ASX stocks on that factor.', italic=True, size=9)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — ANOMALY FLAGS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 10 — Anomaly Flags  (market.anomalies)', level=1, color=RED_DARK)
add_separator()
add_para('Engine: compute/engine/anomaly_detect.py  |  Source: screener.universe  |  Evaluated: hourly (APScheduler)  |  Table: market.anomalies', italic=True)
add_para('Schema: id (UUID), asx_code, flag_type, description, severity, is_active (bool), detected_at, resolved_at  |  Unique constraint: (asx_code, flag_type)')
doc.add_paragraph()
add_table(
    headers=['Flag Type', 'Condition (from screener.universe)', 'Severity', 'Expires'],
    rows=[
        ['volume_spike',  'volume > volume_avg_20d × 3.0',                              'HIGH if >5×, MEDIUM if 3–5×',  '1 day'],
        ['price_gap',     'ABS(open - prev_close) / prev_close > 0.03  (3% gap)',       'HIGH if >5%, MEDIUM if 3–5%',  '3 days'],
        ['reversal',      'Bullish: return_1d > +5% AND prev_return_1d < -3%\nBearish: return_1d < -5% AND prev_return_1d > +3%', 'MEDIUM', '5 days'],
        ['breakout',      'close > 52w high (breakout up)\nclose < 52w low (breakdown)', 'HIGH',                        '10 days'],
        ['rsi_extreme',   'rsi_14 > 80 (overbought) OR rsi_14 < 20 (oversold)',         'MEDIUM',                       'Until RSI re-enters 30–70'],
        ['trend_break',   'Golden cross: sma_50 crosses above sma_200 today\nDeath cross: sma_50 crosses below sma_200 today', 'HIGH', '20 days'],
        ['new_high_low',  'close = new 20d / 52w / all-time high or low',               'MEDIUM',                       '1 day'],
    ],
    col_widths=[1.3, 2.7, 1.7, 1.9],
    hdr_fill='8B0000',
)
add_para('Additional inputs read by anomaly_detect.py: return_3m, earnings_growth_1y, pe_ratio, revenue_growth_1y, short_interest_chg_1w, piotroski_f_score, short_pct, grossed_up_yield, status', italic=True, size=9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — COMPLETE METRICS SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 11 — Complete Metrics Count Summary', level=1, color=BLUE_DARK)
add_separator()
add_table(
    headers=['Table / Layer', 'Schema', 'Type', 'Approx. Metrics', 'Frequency'],
    rows=[
        ['market.daily_prices',        'market',     'Raw download',     '12 fields',           'Daily'],
        ['financials.annual_pnl',      'financials', 'Raw download',     '37 fields',           'Weekly'],
        ['financials.annual_balance_sheet','financials','Raw download',  '36 fields',           'Weekly'],
        ['financials.annual_cashflow', 'financials', 'Raw download',     '22 fields',           'Weekly'],
        ['financials.half_year_pnl',   'financials', 'Raw download',     '20 fields',           'Weekly'],
        ['market.dividends',           'market',     'Raw download',     '10 fields',           'Daily + Weekly'],
        ['market.short_positions',     'market',     'Raw download',     '5 fields',            'Daily (ASIC)'],
        ['market.announcements',       'market',     'Raw download',     '8 fields',            'Every 2 hours'],
        ['market.daily_metrics',       'market',     'Calculated',       '90+ metrics',         'Daily'],
        ['market.weekly_metrics',      'market',     'Calculated',       '43 metrics',          'Weekly'],
        ['market.monthly_metrics',     'market',     'Calculated',       '26 metrics',          'Monthly'],
        ['market.quarterly_metrics',   'market',     'Derived',          '27 metrics',          'Quarterly'],
        ['market.halfyearly_metrics',  'market',     'Derived',          '25 metrics',          'Half-yearly'],
        ['market.yearly_metrics',      'market',     'Derived+Calculated','150+ metrics',       'Annual'],
        ['market.period_metrics',      'market',     'Calculated',       '21 metrics',          'Daily'],
        ['market.anomalies',           'market',     'Detected',         '7 flag types',        'Hourly'],
        ['screener.universe',          'screener',   'All combined',     '458+ columns',        'Nightly'],
    ],
    col_widths=[2.2, 1.0, 1.2, 1.3, 1.9],
    hdr_fill='003E6B',
)

add_heading('Metrics Classification Summary', level=2, color=BLUE_MID)
add_table(
    headers=['Category', 'Count', 'Where Stored'],
    rows=[
        ['Raw Downloaded — Price (EODHD)',    '12',   'market.daily_prices'],
        ['Raw Downloaded — P&L',              '37',   'financials.annual_pnl'],
        ['Raw Downloaded — Balance Sheet',    '36',   'financials.annual_balance_sheet'],
        ['Raw Downloaded — Cash Flow',        '22',   'financials.annual_cashflow'],
        ['Raw Downloaded — Half-Year P&L',    '20',   'financials.half_year_pnl'],
        ['Raw Downloaded — Dividends',        '10',   'market.dividends'],
        ['Raw Downloaded — Short Positions',  '5',    'market.short_positions'],
        ['Raw Downloaded — Announcements',    '8',    'market.announcements'],
        ['Technical — Daily',                 '90+',  'market.daily_metrics'],
        ['Technical — Weekly',                '43',   'market.weekly_metrics'],
        ['Technical — Monthly',               '26',   'market.monthly_metrics'],
        ['Fundamental — Quarterly (Derived)', '27',   'market.quarterly_metrics'],
        ['Fundamental — Half-Yearly (Derived)','25',  'market.halfyearly_metrics'],
        ['Fundamental Ratios — Annual',       '150+', 'market.yearly_metrics'],
        ['Period Price Ranges',               '21',   'market.period_metrics'],
        ['Composite Scores',                  '6',    'screener.universe'],
        ['Anomaly Flags',                     '7',    'market.anomalies'],
        ['TOTAL (unique metrics)',            '550+', 'Across all tables'],
    ],
    col_widths=[3.0, 0.8, 3.8],
    hdr_fill='003E6B',
)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUTPUT_PATH)
print(f'Saved: {OUTPUT_PATH}')
