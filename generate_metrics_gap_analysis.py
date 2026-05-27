"""
ASX Screener — Metrics Coverage & Gap Analysis Word Document
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\Users\Dell\My Claude\Code\ASX Screener\ASX_Screener_Metrics_Coverage_Analysis.docx"

doc = Document()
section = doc.sections[0]
section.top_margin    = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin   = Cm(1.8)
section.right_margin  = Cm(1.8)

BLUE_DARK = (0, 62, 107);  BLUE_MID = (0, 102, 160)
GREEN_OK  = (0, 110, 55);  AMBER = (160, 90, 0)
RED_MISS  = (150, 20, 20); GREY  = (80, 80, 80)
WHITE     = (255, 255, 255)

HEX_GREEN = 'C6EFCE'; HEX_AMBER = 'FFEB9C'; HEX_RED = 'FFC7CE'
HEX_HDR   = '003E6B'; HEX_S1 = 'F0F4FA'; HEX_S2 = 'FFFFFF'

def rgb(t): return RGBColor(*t)

def tbl_shd(cell, fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), fill); tcPr.append(shd)

def add_heading(text, level=1, color=BLUE_DARK):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = rgb(color); run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(8 if level <= 2 else 4)
    p.paragraph_format.space_after  = Pt(3)

def add_para(text='', bold=False, size=10, color=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(size)
    r.font.bold=bold; r.font.italic=italic
    if color: r.font.color.rgb = rgb(color)
    return p

def add_sep():
    p = doc.add_paragraph('─'*140)
    for r in p.runs:
        r.font.name='Courier New'; r.font.size=Pt(5)
        r.font.color.rgb=rgb((180,180,180))
    p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)

def add_legend():
    p = doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
    for icon, col, desc in [
        ('✅ Present', GREEN_OK, '  = Metric exists as a dedicated column   '),
        ('⚠️ Partial',  AMBER,   '  = Component data exists; ratio not pre-computed   '),
        ('❌ Missing',  RED_MISS, '  = Not currently in the platform'),
    ]:
        r1 = p.add_run(icon); r1.font.bold=True; r1.font.size=Pt(9)
        r1.font.name='Calibri'; r1.font.color.rgb=rgb(col)
        r2 = p.add_run(desc); r2.font.size=Pt(9); r2.font.name='Calibri'

# Status → fill colour
S_FILL = {'present': HEX_GREEN, 'partial': HEX_AMBER, 'missing': HEX_RED}

def mtable(headers, rows, cw):
    """
    headers : list of str (N items)
    rows    : list of (status, col0, col1, ...) where cols = N items
    cw      : list of float widths (N items)
    Row colour determined by status; status is NOT rendered as a column.
    """
    assert len(headers) == len(cw), f"headers({len(headers)}) != cw({len(cw)})"
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text=h
        r = c.paragraphs[0].runs[0] if c.paragraphs[0].runs else c.paragraphs[0].add_run(h)
        r.font.bold=True; r.font.size=Pt(8.5); r.font.name='Calibri'
        r.font.color.rgb=rgb(WHITE); tbl_shd(c, HEX_HDR)
    for row in rows:
        status = row[0]; vals = row[1:]
        assert len(vals) == len(headers), f"row cols({len(vals)}) != headers({len(headers)})"
        cells = t.add_row().cells; fill = S_FILL[status]
        for i, v in enumerate(vals):
            cells[i].text=str(v)
            r = cells[i].paragraphs[0].runs[0] if cells[i].paragraphs[0].runs else cells[i].paragraphs[0].add_run(str(v))
            r.font.size=Pt(8.5); r.font.name='Calibri'; tbl_shd(cells[i], fill)
    for row in t.rows:
        for i, w in enumerate(cw):
            row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(3)

# ── COVER ─────────────────────────────────────────────────────────────────────
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ASX SCREENER')
r.font.name='Calibri'; r.font.size=Pt(28); r.font.bold=True; r.font.color.rgb=rgb(BLUE_DARK)

p2 = doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Metrics Coverage & Gap Analysis')
r2.font.name='Calibri'; r2.font.size=Pt(16); r2.font.color.rgb=rgb(BLUE_MID)

doc.add_paragraph()
p3 = doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Mapping all 11 metric categories against the actual database schema\n'
                 'Identifying what is present, partial, and missing')
r3.font.name='Calibri'; r3.font.size=Pt(11); r3.font.color.rgb=rgb(GREY)

doc.add_paragraph()
p4 = doc.add_paragraph(); p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('✅ Present: 67 metrics   ⚠️ Partial: 5 metrics   ❌ Missing: 13 metrics'
                 '   |   Overall Coverage: 79%')
r4.font.name='Calibri'; r4.font.size=Pt(12); r4.font.bold=True

doc.add_page_break()

# ── EXECUTIVE SUMMARY ─────────────────────────────────────────────────────────
add_heading('Executive Summary', level=1)
add_legend()
add_heading('Coverage by Category', level=2, color=BLUE_MID)
mtable(
    headers=['Category', '✅ Present', '⚠️ Partial', '❌ Missing', 'Coverage'],
    rows=[
        ('present','1. Fundamental Metrics',      '7','0','0','100%'),
        ('present','2. Profitability Metrics',     '7','0','0','100%'),
        ('present','3. Valuation Metrics',         '9','0','0','100%'),
        ('present','4. Growth Metrics',            '8','0','0','100%'),
        ('present','5. Financial Health Metrics',  '8','0','0','100%'),
        ('partial','6. Cash Flow Quality',         '3','1','1','70%'),
        ('partial','7. Technical Indicators',      '8','1','0','94%'),
        ('partial','8. Risk Metrics',              '3','1','4','44%'),
        ('partial','9. Dividend Metrics',          '5','0','1','83%'),
        ('partial','10. Momentum Metrics',         '6','0','1','86%'),
        ('missing','11. Quality & Moat',           '3','2','6','36%'),
        ('partial','TOTAL',                        '67','5','13','79%'),
    ],
    cw=[3.2, 0.9, 0.9, 0.9, 0.9],
)
add_para('Categories 1–5 are 100% complete. Categories 6–10 have minor computable gaps only. '
         'Category 11 (Quality & Moat) has 6 qualitative indicators that cannot be automated.',
         italic=True, size=9)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — FUNDAMENTAL
# ══════════════════════════════════════════════════════════════════════════════
add_heading('1.  Fundamental Metrics  ✅ 7/7 Present', level=1, color=GREEN_OK)
add_sep(); add_legend()
mtable(
    headers=['Metric', 'DB Column', 'Table(s)', 'Notes'],
    rows=[
        ('present','✅ Revenue',              'revenue',                     'financials.annual_pnl\nscreener.universe',         'Trailing 12M + all historical years'),
        ('present','✅ Revenue Growth',        'revenue_growth_1y',           'market.yearly_metrics\nscreener.universe',         'YoY. Also CAGR 3Y/5Y/7Y/10Y'),
        ('present','✅ Net Profit',            'net_profit',                  'financials.annual_pnl\nscreener.universe',         'After minority interest'),
        ('present','✅ EPS',                  'eps / eps_diluted',            'financials.annual_pnl\nmarket.yearly_metrics',     'Basic and diluted both present'),
        ('present','✅ Book Value per Share',  'bvps',                        'market.yearly_metrics\nscreener.universe',         'Tangible BVPS (tbvps) also available'),
        ('present','✅ Market Capitalisation', 'market_cap',                  'market.yearly_metrics\nscreener.universe',         'Recomputed daily at latest price'),
        ('present','✅ Sector / Industry',     'sector / industry\ngics_sector / gics_industry',
                                               'market.securities\nscreener.universe',              '4-level GICS classification. Always filter within sector.'),
    ],
    cw=[1.9, 1.8, 2.0, 2.3],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — PROFITABILITY
# ══════════════════════════════════════════════════════════════════════════════
add_heading('2.  Profitability Metrics  ✅ 7/7 Present', level=1, color=GREEN_OK)
add_sep(); add_legend()
mtable(
    headers=['Metric', 'DB Column', 'Table(s)', 'Good Range', 'Notes'],
    rows=[
        ('present','✅ Gross Margin',     'gross_margin',  'market.yearly_metrics\nscreener.universe', 'Higher than peers',  'avg_gross_margin_3y / 5y also available'),
        ('present','✅ EBITDA Margin',    'ebitda_margin', 'market.yearly_metrics\nscreener.universe', '15%+',              'avg_ebitda_margin_3y / 5y also available'),
        ('present','✅ EBIT Margin',      'ebit_margin',   'market.yearly_metrics\nscreener.universe', '10%+',              'avg_operating_margin_3y / 5y / 10y'),
        ('present','✅ Net Profit Margin','net_margin',    'market.yearly_metrics\nscreener.universe', '8–10%+',            'avg_net_margin_3y / 5y / 10y'),
        ('present','✅ ROE',              'roe',           'market.yearly_metrics\nscreener.universe', '15%+',              'avg_roe_3y / 5y / 7y / 10y. Also roae, roaa.'),
        ('present','✅ ROA',              'roa',           'market.yearly_metrics\nscreener.universe', '5%+',               'avg_roa_3y / 5y'),
        ('present','✅ ROIC',             'roic',          'market.yearly_metrics\nscreener.universe', '12–15%+',           'avg_roic_3y / 5y. croic (cash ROIC) also available. Best quality indicator.'),
    ],
    cw=[1.7, 1.5, 1.8, 1.0, 2.0],
)
add_para('Extra profitability metrics also present: roce, roae, roaa, ocf_margin, fcf_margin, pretax_margin, tax_rate_effective.', italic=True, size=9)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — VALUATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('3.  Valuation Metrics  ✅ 9/9 Present', level=1, color=GREEN_OK)
add_sep(); add_legend()
mtable(
    headers=['Metric', 'DB Column', 'Table(s)', 'Notes'],
    rows=[
        ('present','✅ P/E Ratio',         'pe_ratio',         'market.yearly_metrics\nscreener.universe','Trailing. pe_5y_avg also stored.'),
        ('present','✅ Forward P/E',        'pe_ratio_forward', 'screener.universe',                       'Migration 011. Populated from EODHD forecast EPS.'),
        ('present','✅ PEG Ratio',          'peg_ratio',        'market.yearly_metrics\nscreener.universe','pe_ratio / eps_growth_1y. <1 attractive.'),
        ('present','✅ Price to Book',      'pb_ratio',         'market.yearly_metrics\nscreener.universe','price / bvps. Useful for banks/resources.'),
        ('present','✅ EV/EBITDA',          'ev_ebitda',        'market.yearly_metrics\nscreener.universe','Capital-structure-neutral. Preferred for comparisons.'),
        ('present','✅ EV/EBIT',            'ev_ebit',          'market.yearly_metrics\nscreener.universe','Useful for profitable firms.'),
        ('present','✅ Price to Sales',     'ps_ratio',         'market.yearly_metrics\nscreener.universe','market_cap / revenue. Good for early-growth.'),
        ('present','✅ FCF Yield',          'fcf_yield',        'market.yearly_metrics\nscreener.universe','fcf / market_cap. 4–6%+ good.'),
        ('present','✅ Dividend Yield',     'dividend_yield',   'market.yearly_metrics\nscreener.universe','Standard + grossed_up_yield (includes franking credit).'),
    ],
    cw=[1.9, 1.7, 2.0, 2.4],
)
add_para('Extra valuation metrics: pcf_ratio, p_fcf_ratio, ev_revenue, ev_fcf, earnings_yield, graham_number, ncavps.', italic=True, size=9)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — GROWTH
# ══════════════════════════════════════════════════════════════════════════════
add_heading('4.  Growth Metrics  ✅ 8/8 Present', level=1, color=GREEN_OK)
add_sep(); add_legend()
mtable(
    headers=['Metric', 'DB Column(s)', 'Table(s)', 'Notes'],
    rows=[
        ('present','✅ Revenue Growth YoY',   'revenue_growth_1y',                          'market.yearly_metrics\nscreener.universe','Year-on-year.'),
        ('present','✅ Revenue CAGR 3Y / 5Y', 'revenue_cagr_3y\nrevenue_cagr_5y',           'market.yearly_metrics\nscreener.universe','Also 7Y, 10Y, median_5y, median_10y.'),
        ('present','✅ EPS Growth YoY',        'eps_growth_1y',                              'market.yearly_metrics\nscreener.universe','Year-on-year.'),
        ('present','✅ EPS CAGR 3Y / 5Y',     'eps_cagr_3y\neps_cagr_5y',                  'market.yearly_metrics\nscreener.universe','Also 7Y, 10Y. avg_eps_growth_3y/5y/10y.'),
        ('present','✅ EBITDA Growth',         'ebitda_growth_1y\nebitda_cagr_3y/5y/7y/10y','market.yearly_metrics\nscreener.universe','YoY + CAGR across 4 timeframes.'),
        ('present','✅ FCF Growth',            'fcf_growth_1y\nfcf_cagr_3y\nfcf_cagr_5y',  'market.yearly_metrics\nscreener.universe','YoY + 3Y + 5Y CAGR.'),
        ('present','✅ Book Value Growth',     'bvps_growth_1y\nbvps_cagr_3y\nbvps_cagr_5y','market.yearly_metrics\nscreener.universe','YoY + 3Y + 5Y CAGR.'),
        ('present','✅ Dividend Growth',       'dividend_cagr_3y\ndividend_cagr_5y',         'market.yearly_metrics\nscreener.universe','DPS CAGR 3Y and 5Y.'),
    ],
    cw=[2.0, 2.2, 1.8, 2.0],
)
add_para('Extra growth metrics: gross_profit_growth_1y, net_income_cagr_3y/5y/7y/10y, ocf_growth_1y, price_cagr_1y/3y/5y/7y/10y.', italic=True, size=9)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — FINANCIAL HEALTH
# ══════════════════════════════════════════════════════════════════════════════
add_heading('5.  Financial Health Metrics  ✅ 8/8 Present', level=1, color=GREEN_OK)
add_sep(); add_legend()
mtable(
    headers=['Metric', 'DB Column(s)', 'Table(s)', 'Good Range', 'Notes'],
    rows=[
        ('present','✅ Debt to Equity',       'debt_to_equity',                         'market.yearly_metrics\nscreener.universe','<0.5 conservative','debt_to_assets, net_debt_to_equity also present.'),
        ('present','✅ Net Debt / EBITDA',    'net_debt_to_ebitda',                     'market.yearly_metrics\nscreener.universe','<2 good, <1 strong','debt_to_ebitda (gross) also available.'),
        ('present','✅ Interest Coverage',    'interest_coverage',                      'market.yearly_metrics\nscreener.universe','>5 strong',         'EBIT / interest_expense.'),
        ('present','✅ Current Ratio',        'current_ratio',                          'market.yearly_metrics\nscreener.universe','>1.5 good',          'avg_current_ratio_3y also stored.'),
        ('present','✅ Quick Ratio',          'quick_ratio',                            'market.yearly_metrics\nscreener.universe','>1 good',            'cash_ratio also available.'),
        ('present','✅ Cash Balance',         'cash_equivalents',                       'financials.annual_balance_sheet\nscreener.universe','Higher better','net_debt, working_capital also present.'),
        ('present','✅ Operating Cash Flow',  'cfo\nocf_margin',                        'financials.annual_cashflow\nmarket.yearly_metrics','Positive','ocf_margin=cfo/revenue. ocf_growth_1y also available.'),
        ('present','✅ Free Cash Flow',       'fcf\nfcf_margin\nfcf_yield\nfcf_per_share','financials.annual_cashflow\nmarket.yearly_metrics','Positive + growing','Full FCF suite: amount, margin, yield, per share.'),
    ],
    cw=[1.7, 1.9, 1.8, 1.0, 1.6],
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — CASH FLOW QUALITY
# ══════════════════════════════════════════════════════════════════════════════
add_heading('6.  Cash Flow Quality Metrics  ⚠️ 3✅ 1⚠️ 1❌  (70% coverage)', level=1, color=AMBER)
add_sep(); add_legend()
mtable(
    headers=['Metric', 'DB Column', 'Table(s)', 'Notes'],
    rows=[
        ('partial','⚠️ OCF / Net Profit\n(Cash Earnings Quality)',
         'fcf_to_net_income exists\nocf_to_net_profit MISSING',
         'screener.universe (partial)',
         'DATA AVAILABLE. cfo and net_profit both exist. ocf_to_net_profit ratio '
         'not pre-computed but trivially addable to market.yearly_metrics. '
         'Related: fcf_to_net_income is already in screener.universe.'),
        ('present','✅ FCF Margin (FCF / Revenue)',
         'fcf_margin',
         'market.yearly_metrics\nscreener.universe',
         'Fully implemented. avg_fcf_margin_3y / 5y also stored.'),
        ('present','✅ Capex / Revenue (Capital Intensity)',
         'capex_intensity',
         'market.yearly_metrics\nscreener.universe',
         'capex / revenue. Also raw capex in financials.annual_cashflow.'),
        ('missing','❌ FCF Positive Years\n(Consecutive count)',
         'NOT present',
         'Nowhere in schema',
         'DATA AVAILABLE. fcf stored per year in financials.annual_cashflow. '
         'Easy to add as rolling COUNT of consecutive years with fcf > 0 to market.yearly_metrics.'),
        ('present','✅ Cash Conversion Cycle',
         'cash_conversion_cycle',
         'market.yearly_metrics\nscreener.universe',
         'CCC = receivables_days + inventory_days − payables_days. '
         'All three components also stored separately.'),
    ],
    cw=[2.2, 2.0, 1.8, 2.0],
)
add_heading('Action Items', level=3, color=RED_MISS)
add_para('1.  Add ocf_to_net_profit = cfo / net_profit  →  market.yearly_metrics + screener.universe.   (1 column, all data available, 1-day effort)', bold=True)
add_para('2.  Add fcf_positive_years = COUNT(consecutive years with fcf > 0)  →  market.yearly_metrics.   (1 column, all data available, 1-day effort)', bold=True)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — TECHNICAL
# ══════════════════════════════════════════════════════════════════════════════
add_heading('7.  Technical Indicators  ⚠️ 8✅ 1⚠️  (94% coverage)', level=1, color=BLUE_MID)
add_sep(); add_legend()
mtable(
    headers=['Indicator', 'DB Column(s)', 'Table(s)', 'Notes'],
    rows=[
        ('present','✅ 50-Day Moving Average',
         'sma_50\ndma50_ratio\nabove_sma50',
         'market.daily_metrics\nscreener.universe',
         'Value + ratio vs price + binary flag. Also sma_5/10/20/100/200 all present.'),
        ('present','✅ 200-Day Moving Average',
         'sma_200\ndma200_ratio\nabove_sma200',
         'market.daily_metrics\nscreener.universe',
         'Full suite of EMA also: ema_9/12/20/26/50/200.'),
        ('present','✅ 50DMA > 200DMA (Cross signals)',
         'golden_cross\ndeath_cross\nmacd_bullish_cross\nmacd_bearish_cross',
         'market.daily_metrics\nscreener.universe',
         'Binary flags for both daily MA cross and MACD cross.'),
        ('present','✅ RSI',
         'rsi_14\nrsi_7\nrsi_21\nrsi_overbought\nrsi_oversold',
         'market.daily_metrics\nscreener.universe',
         '3 periods (7/14/21). Overbought (>70) and oversold (<30) binary flags.'),
        ('present','✅ MACD',
         'macd_line\nmacd_signal\nmacd_hist\nmacd_bullish_cross\nmacd_bearish_cross',
         'market.daily_metrics\nscreener.universe',
         'Full MACD suite. Weekly + Monthly MACD also in respective tables.'),
        ('present','✅ Volume Trend',
         'relative_volume\nvolume_avg_20d\nobv\ncmf_20\nmfi_14',
         'market.daily_metrics\nscreener.universe',
         'relative_volume=vol/20d_avg. OBV, CMF, MFI also stored. Volume spike anomaly active.'),
        ('present','✅ Relative Strength vs Index',
         'relative_strength_xjo\nmomentum_score',
         'market.monthly_metrics\nscreener.universe',
         'return_12m minus ASX 200 return_12m. Momentum percentile rank 0–100 in universe.'),
        ('present','✅ Price near 52-Week High',
         'pct_from_52w_high\nnew_52w_high\nhigh_52w',
         'market.daily_metrics\nscreener.universe',
         '% below 52w high + new high flag + absolute 52w high level. pct_from_ath also available.'),
        ('partial','⚠️ Price above VWAP',
         'vwap (value stored)\nno above_vwap binary flag',
         'market.daily_metrics',
         'VWAP is computed and stored. A binary above_vwap flag is NOT pre-computed. '
         'Trivially addable: CASE WHEN close > vwap THEN 1 ELSE 0 END.'),
    ],
    cw=[2.0, 2.1, 1.8, 2.1],
)
add_heading('Action Item', level=3, color=AMBER)
add_para('Add above_vwap = CASE WHEN close > vwap THEN 1 ELSE 0 END  →  market.daily_metrics + screener.universe.   (1 column, data available, 1-day effort)', bold=True)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — RISK
# ══════════════════════════════════════════════════════════════════════════════
add_heading('8.  Risk Metrics  ⚠️ 3✅ 1⚠️ 4❌  (44% coverage)', level=1, color=RED_MISS)
add_sep(); add_legend()
mtable(
    headers=['Metric', 'DB Column(s)', 'Table(s)', 'Notes'],
    rows=[
        ('present','✅ Beta',
         'beta_1y\nbeta_3y\nbeta_5y',
         'market.yearly_metrics\nscreener.universe',
         'Three timeframes vs ASX 200. alpha_1y / alpha_3y also present. '
         '<1 = lower market risk, >1 = higher.'),
        ('present','✅ Share Price Drawdown',
         'pct_from_ath\nmax_drawdown_1y\nmax_drawdown_3y\ncalmar_ratio',
         'market.daily_metrics\nmarket.yearly_metrics\nscreener.universe',
         'Daily drawdown from ATH + max drawdown over 1Y and 3Y windows. '
         'calmar_ratio = return/max_drawdown. Volatility_1y/3y also present.'),
        ('present','✅ Debt Risk',
         'debt_to_equity\nnet_debt_to_ebitda\ninterest_coverage\ndebt_to_assets',
         'market.yearly_metrics\nscreener.universe',
         'Comprehensive debt suite. Also lt_debt_to_capital, net_debt_per_share.'),
        ('partial','⚠️ Liquidity Risk (trading)',
         'volume_avg_20d\nvolume_avg_52w\nrelative_volume',
         'market.daily_metrics\nscreener.universe',
         'Volume metrics let you filter illiquid stocks (e.g. avg_vol < $500K/day). '
         'No single "liquidity_risk_score" but components all present.'),
        ('missing','❌ Earnings Volatility\n(EPS stability)',
         'NOT present',
         'Nowhere in schema',
         'DATA AVAILABLE. eps stored per year in market.yearly_metrics. '
         'Add eps_volatility_5y = STDDEV(eps,5y)/AVG(eps,5y) = coefficient of variation. '
         'Low = stable earnings. 1-day effort.'),
        ('missing','❌ Dilution Risk\n(Share count change)',
         'NOT present',
         'Nowhere in schema',
         'DATA AVAILABLE. shares_outstanding in financials.annual_balance_sheet per year. '
         'Add shares_dilution_3y = (shares_now − shares_3y_ago)/shares_3y_ago. '
         'Repeated positive = dilutive. 1-day effort.'),
        ('missing','❌ Revenue Concentration\n(Customer/product risk)',
         'NOT present',
         'Requires new data',
         'Requires segment-level revenue data from annual reports. '
         'Not available via EODHD. Would need manual tagging or premium data provider.'),
        ('missing','❌ Commodity/Currency Exposure',
         'NOT present',
         'Requires external data',
         'Partially inferred from GICS sector (Materials, Energy). '
         'For specifics: requires company disclosure parsing or manual classification tags.'),
    ],
    cw=[2.1, 2.0, 1.7, 2.2],
)
add_heading('Action Items', level=3, color=RED_MISS)
add_para('1.  Add eps_volatility_5y = STDDEV(eps,5y)/AVG(eps,5y)  →  market.yearly_metrics + screener.universe.   (data available)', bold=True)
add_para('2.  Add shares_dilution_3y = (shares_now − shares_3y_ago)/shares_3y_ago  →  market.yearly_metrics + screener.universe.   (data available)', bold=True)
add_para('3.  Revenue Concentration and Commodity/Currency Exposure require new data sources. Consider adding commodity_exposure and currency_exposure text tags to market.securities as admin-maintained fields.', bold=True)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — DIVIDEND
# ══════════════════════════════════════════════════════════════════════════════
add_heading('9.  Dividend Metrics  ⚠️ 5✅ 1❌  (83% coverage)', level=1, color=BLUE_MID)
add_sep(); add_legend()
mtable(
    headers=['Metric', 'DB Column(s)', 'Table(s)', 'Notes'],
    rows=[
        ('present','✅ Dividend Yield',
         'dividend_yield\nfranked_yield (grossed-up)',
         'market.yearly_metrics\nscreener.universe',
         'Standard yield + grossed-up yield (including franking credit value). 3–6% sustainable.'),
        ('present','✅ Payout Ratio (Div/EPS)',
         'payout_ratio',
         'market.yearly_metrics\nscreener.universe',
         'dividends / eps. 40–70% healthy. >100% = paying more than earned.'),
        ('missing','❌ FCF Payout Ratio (Div/FCF)',
         'NOT present',
         'Nowhere in schema',
         'DATA AVAILABLE. dividends_paid (financials.annual_cashflow) and fcf both exist. '
         'Add fcf_payout_ratio = dividends_paid / fcf. '
         'Preferred over earnings payout as FCF is harder to manipulate. 1-day effort.'),
        ('present','✅ Dividend Growth',
         'dividend_cagr_3y\ndividend_cagr_5y',
         'market.yearly_metrics\nscreener.universe',
         'DPS CAGR over 3 and 5 years.'),
        ('present','✅ Dividend History (Consecutive years)',
         'dividend_consecutive_yrs',
         'market.yearly_metrics\nscreener.universe',
         'Count of consecutive years with dividend payment. 5+ years = reliable.'),
        ('present','✅ Franking Credits (Australian-specific)',
         'franking_pct\nfranked_yield\ndps_grossed_up\nfranking_credit_per_share',
         'market.yearly_metrics\nscreener.universe',
         'Full franking suite. Formula: credit = dps × (franking%/100) × (0.30/0.70). '
         'Grossed-up yield critical for Australian investors comparing pre-tax returns.'),
    ],
    cw=[2.2, 2.0, 1.7, 2.1],
)
add_heading('Action Item', level=3, color=AMBER)
add_para('Add fcf_payout_ratio = dividends_paid / fcf  →  market.yearly_metrics + screener.universe.   (data available, 1-day effort)', bold=True)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — MOMENTUM
# ══════════════════════════════════════════════════════════════════════════════
add_heading('10.  Momentum Metrics  ⚠️ 6✅ 1❌  (86% coverage)', level=1, color=BLUE_MID)
add_sep(); add_legend()
mtable(
    headers=['Metric', 'DB Column(s)', 'Table(s)', 'Notes'],
    rows=[
        ('present','✅ 3-Month Price Return',
         'return_3m',
         'market.daily_metrics\nscreener.universe',
         'Trading-day exact: 63 sessions.'),
        ('present','✅ 6-Month Price Return',
         'return_6m',
         'market.daily_metrics\nscreener.universe',
         'Trading-day exact: 126 sessions.'),
        ('present','✅ 12-Month Price Return',
         'return_1y',
         'market.daily_metrics\nscreener.universe',
         'Trading-day exact: 252 sessions. Also return_3y/5y/7y/10y/15y available.'),
        ('present','✅ Relative Strength Rank',
         'relative_strength_xjo\nmomentum_score',
         'market.monthly_metrics\nscreener.universe',
         'relative_strength_xjo = return_12m minus ASX 200 return. '
         'momentum_score = percentile rank 0–100 within all ASX stocks.'),
        ('missing','❌ Earnings Revision Trend\n(Analyst upgrades/downgrades)',
         'NOT present',
         'Requires new data source',
         'Requires analyst consensus data (Bloomberg, Refinitiv, FactSet). '
         'EODHD has /analyst-ratings/{ticker} endpoint (buy/sell/hold counts + consensus EPS). '
         'Integrating this would also fully populate pe_ratio_forward. Recommended.'),
        ('present','✅ Volume Breakout',
         'relative_volume\nhas_volume_spike (anomaly)',
         'market.daily_metrics\nscreener.universe\nmarket.anomalies',
         'relative_volume>2 = elevated. volume_spike flag at 3× average. '
         'has_volume_spike in screener.universe.'),
        ('present','✅ 52-Week High Proximity',
         'pct_from_52w_high\nnew_52w_high\nhigh_52w',
         'market.daily_metrics\nscreener.universe',
         '% below 52w high + new high flag + absolute level. '
         'Within 10–15% of 52w high = strong trend signal.'),
    ],
    cw=[2.2, 2.0, 1.7, 2.1],
)
add_heading('Action Item', level=3, color=RED_MISS)
add_para('Earnings Revision Trend: Integrate EODHD /analyst-ratings/{ticker} endpoint. '
         'Returns buy/sell/hold counts and consensus EPS estimates. '
         'Would enable eps_revision_trend column and fully populate pe_ratio_forward. '
         'Recommended as the highest-value data source addition.', bold=True)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — QUALITY & MOAT
# ══════════════════════════════════════════════════════════════════════════════
add_heading('11.  Quality & Moat Indicators  ⚠️ 3✅ 2⚠️ 6❌  (36% coverage)', level=1, color=RED_MISS)
add_sep(); add_legend()
mtable(
    headers=['Indicator', 'DB Column / Proxy', 'Table(s)', 'Notes'],
    rows=[
        ('present','✅ ROIC Consistency',
         'roic\navg_roic_3y\navg_roic_5y',
         'market.yearly_metrics\nscreener.universe',
         'ROIC + 3Y and 5Y averages. Sustained ROIC >12–15% = durable competitive advantage.'),
        ('present','✅ Gross Margin Stability',
         'gross_margin\navg_gross_margin_3y\navg_gross_margin_5y',
         'market.yearly_metrics\nscreener.universe',
         'Gross margin + 3Y and 5Y averages. Stable or improving = pricing power / moat.'),
        ('present','✅ Piotroski F-Score\n(Composite quality proxy)',
         'piotroski_f_score',
         'market.yearly_metrics\nscreener.universe',
         '9-point score: 3×profitability + 3×leverage + 3×efficiency signals. '
         'Score 8–9 = high financial quality. Altman Z + Beneish M also available.'),
        ('partial','⚠️ Asset-Light Model',
         'capex_intensity\nfcf_margin\nasset_turnover',
         'market.yearly_metrics\nscreener.universe',
         'No single "asset_light" flag. Proxy: low capex_intensity + high fcf_margin + '
         'high asset_turnover = asset-light characteristics. Can add a composite.'),
        ('partial','⚠️ Management Quality\n(Capital Allocation)',
         'avg_roic_5y / avg_roe_5y\nprice_cagr_5y / price_cagr_10y',
         'market.yearly_metrics\nscreener.universe',
         'Proxied by long-run ROIC/ROE consistency and long-run price CAGR. '
         'Insider holding % was in v009 schema but removed in v011. '
         'Recommend re-adding from EODHD /insider-transactions.'),
        ('missing','❌ Recurring Revenue',
         'NOT present',
         'Manual classification needed',
         'No revenue type classification. Requires business model tagging '
         '(SaaS/subscription/transactional/project). Cannot be derived from financials alone.'),
        ('missing','❌ Low Churn',
         'NOT present',
         'Company-specific data',
         'Customer retention not in any public financial filing. '
         'Would need scraping of company presentations or industry databases.'),
        ('missing','❌ Market Share',
         'NOT present',
         'Requires industry data',
         'Cannot be derived from financial statements alone. '
         'Would need external industry revenue data by sector (IBISWorld, Statista, etc.).'),
        ('missing','❌ Brand Strength',
         'NOT present — qualitative',
         'N/A',
         'Cannot be quantified from financial data. '
         'Partial proxy: sustained high gross margin and premium ps_ratio over 5+ years.'),
        ('missing','❌ Network Effects',
         'NOT present — qualitative',
         'N/A',
         'Cannot be quantified. Partial proxy: high revenue growth with low capex_intensity '
         '(digital scaling without proportional cost).'),
        ('missing','❌ Switching Costs',
         'NOT present — qualitative',
         'N/A',
         'Cannot be quantified. Partial proxy: stable customer base = stable/growing revenue '
         'with low revenue volatility (eps_volatility_5y once added).'),
    ],
    cw=[2.0, 2.2, 1.7, 2.1],
)
add_heading('Action Items', level=3, color=RED_MISS)
add_para('1.  Re-add insider_holding_pct and insider_holding_change_1y to screener.universe '
         '(was in v009, removed in v011). Source: EODHD /insider-transactions endpoint.', bold=True)
add_para('2.  Consider adding business_model_tag text column to market.securities for manual '
         'classification (SaaS, Subscription, Mining, REIT, etc.) to partially enable '
         '"recurring revenue" and "asset-light" filtering.', bold=True)
add_para('3.  Qualitative moat indicators (brand, network effects, switching costs) cannot be '
         'automated from financial data. Consider a moat_rating or analyst_notes field in '
         'market.securities for admin-maintained qualitative signals.', bold=True)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — GAPS SUMMARY & ACTION PLAN
# ══════════════════════════════════════════════════════════════════════════════
add_heading('12.  Complete Gap Summary & Recommended Action Plan', level=1, color=BLUE_DARK)
add_sep()

add_heading('12.1  Missing Metrics — Computable Today (All Data Already Available)', level=2, color=GREEN_OK)
add_para('These 7 items require only adding new computed columns to existing tables. No new data source needed.')
mtable(
    headers=['Priority', 'Metric to Add', 'Formula', 'Target Table(s)', 'Effort'],
    rows=[
        ('present','HIGH',   'OCF / Net Profit',          'cfo / net_profit',                             'market.yearly_metrics\nscreener.universe', '½ day'),
        ('present','HIGH',   'FCF Payout Ratio',           'dividends_paid / fcf',                         'market.yearly_metrics\nscreener.universe', '½ day'),
        ('present','HIGH',   'Dilution Risk (3Y)',          '(shares_now − shares_3y_ago) / shares_3y_ago', 'market.yearly_metrics\nscreener.universe', '½ day'),
        ('present','MEDIUM', 'Earnings Volatility',        'STDDEV(eps,5y) / AVG(eps,5y)',                 'market.yearly_metrics\nscreener.universe', '½ day'),
        ('present','MEDIUM', 'FCF Positive Years',         'COUNT(consecutive years with fcf > 0)',         'market.yearly_metrics\nscreener.universe', '1 day'),
        ('present','MEDIUM', 'Price Above VWAP (flag)',    'CASE WHEN close > vwap THEN 1 ELSE 0 END',     'market.daily_metrics\nscreener.universe',   '½ day'),
        ('present','MEDIUM', 'Insider Holding %',          'Re-add from EODHD /insider-transactions',      'screener.universe',                         '2 days'),
    ],
    cw=[0.8, 1.8, 2.5, 1.8, 0.9],
)

add_heading('12.2  Missing Metrics — Require a New Data Source', level=2, color=RED_MISS)
add_para('These 2 items require integrating a new API feed.')
mtable(
    headers=['Priority', 'Metric', 'Recommended Data Source', 'What It Enables', 'Effort'],
    rows=[
        ('missing','HIGH',
         'Earnings Revision Trend\n(Analyst upgrades/downgrades)',
         'EODHD /analyst-ratings/{ticker}\n(buy/hold/sell + consensus EPS)',
         'eps_revision_trend, consensus_eps,\nfully populates pe_ratio_forward',
         '3–5 days'),
        ('missing','MEDIUM',
         'Revenue Concentration\nCommodity / Currency Exposure',
         'Manual admin tags on market.securities\n(commodity_exposure, currency_exposure)',
         'Qualitative risk filters in screener',
         '1 day (data entry)'),
    ],
    cw=[0.8, 2.0, 2.2, 1.9, 0.9],
)

add_heading('12.3  Metrics That Are Inherently Qualitative (Not Automatable)', level=2, color=GREY)
add_para('These 5 indicators require human judgment. Best approach: admin-maintained tags on market.securities.')
mtable(
    headers=['Indicator', 'Best Available Quantitative Proxy', 'Recommended Approach'],
    rows=[
        ('missing','Brand Strength',    'Sustained gross margin >40% + premium ps_ratio over 5Y',           'moat_rating admin field on market.securities'),
        ('missing','Network Effects',   'High revenue growth with low capex_intensity over 5Y',              'business_model_tag field'),
        ('missing','Switching Costs',   'Low revenue volatility (stable recurring earnings)',                'business_model_tag field'),
        ('missing','Low Churn',         'Stable/growing revenue without M&A = organic retention signal',     'Not automatable from public data'),
        ('missing','Market Share',      'Revenue CAGR vs sector revenue growth (requires sector data)',      'Requires premium industry data source'),
    ],
    cw=[1.8, 3.0, 3.0],
)

# ── FINAL SCORECARD ───────────────────────────────────────────────────────────
doc.add_page_break()
add_heading('13.  Final Coverage Scorecard', level=1, color=BLUE_DARK)
add_sep()
mtable(
    headers=['Category', 'Total', '✅ Present', '⚠️ Partial', '❌ Computable', '❌ Needs Source', '❌ Qualitative'],
    rows=[
        ('present','1. Fundamental',         '7',  '7', '0', '0', '0', '0'),
        ('present','2. Profitability',        '7',  '7', '0', '0', '0', '0'),
        ('present','3. Valuation',            '9',  '9', '0', '0', '0', '0'),
        ('present','4. Growth',               '8',  '8', '0', '0', '0', '0'),
        ('present','5. Financial Health',     '8',  '8', '0', '0', '0', '0'),
        ('partial','6. Cash Flow Quality',    '5',  '3', '1', '1', '0', '0'),
        ('partial','7. Technical',            '9',  '8', '1', '0', '0', '0'),
        ('partial','8. Risk',                 '8',  '3', '1', '2', '1', '1'),
        ('partial','9. Dividend',             '6',  '5', '0', '1', '0', '0'),
        ('partial','10. Momentum',            '7',  '6', '0', '0', '1', '0'),
        ('missing','11. Quality & Moat',      '11', '3', '2', '1', '0', '5'),
        ('partial','TOTAL',                   '85', '67','5', '5', '2', '6'),
    ],
    cw=[2.0, 0.7, 0.8, 0.8, 1.0, 1.1, 1.0],
)
add_para('RESULT: 67 of 85 metrics (79%) are fully implemented as dedicated columns.\n'
         '5 more are partially covered (data exists, ratio not pre-computed).\n'
         '5 gaps are computable from existing data — all closable within ~1 week of development.\n'
         '2 require a new data source (analyst consensus — EODHD has this endpoint).\n'
         '6 are inherently qualitative and cannot be derived from financial data alone.',
         bold=True, size=10)

doc.save(OUTPUT_PATH)
print(f'Saved: {OUTPUT_PATH}')
